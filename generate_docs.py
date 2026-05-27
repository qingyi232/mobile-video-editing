# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clear_cell(cell):
    tc = cell._element
    for p in tc.findall(qn('w:p')):
        tc.remove(p)


def add_paragraph_to_cell(cell, text, font_name='宋体', font_size=Pt(12), bold=False):
    p = cell.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    if bold:
        run.font.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return p


def set_cell_content(cell, body_paragraphs, font_size=Pt(12)):
    clear_cell(cell)
    for para in body_paragraphs:
        add_paragraph_to_cell(cell, para, font_size=font_size)


def generate_task_book():
    doc = Document('任务书模板.docx')

    p3 = doc.paragraphs[3]
    for run in p3.runs:
        if '数据科学与大数据技术' in run.text:
            run.text = run.text.replace('数据科学与大数据技术', '软件工程')

    p7 = doc.paragraphs[7]
    for run in p7.runs:
        run.text = ''
    p7.runs[0].text = '完成期限： 2025 年  11月  20日至    2026年  6月  10日'
    p7.runs[0].font.size = Pt(12)

    p9 = doc.paragraphs[9]
    for run in p9.runs:
        run.text = ''
    p9.runs[0].text = '一、题目名称：     移动端短视频智能剪辑APP的开发与优化'
    p9.runs[0].font.size = Pt(12)

    p13 = doc.paragraphs[13]
    for run in p13.runs:
        run.text = ''
    main_content = (
        "    现在短视频越来越火，但很多人想自己剪个视频还是觉得挺麻烦的[1]，专业软件不好上手，"
        "简单的工具又不够智能。本课题打算做一个手机上用的短视频智能剪辑APP，主要包含这些功能："
        "a.基础的视频编辑，像视频导入、片段裁剪拼接、加转场、配背景音乐和人声分离，还有用语音"
        "识别来自动加字幕。b.智能辅助方面，用CNN+LSTM深度学习模型来做场景识别[2]，按照视频"
        "帧的清晰度来挑选比较好的片段[3]，再根据场景识别的结果推荐合适的音乐。c.性能优化这块，调整"
        "FFmpeg的参数让渲染更快[4]，用图像压缩减少内存占用，优化界面渲染减少卡顿。d.还有一些"
        "特色功能，比如提供节日、vlog、教程这些模板，支持横竖屏切换，能导出MP4和AVI格式。"
    )
    p13.runs[0].text = main_content
    p13.runs[0].font.size = Pt(12)

    p14 = doc.paragraphs[14]
    for run in p14.runs:
        run.text = ''
    second_para = (
        "     技术上前端用Flutter[5]做跨平台开发，Android和iOS都能用。后端用Spring Boot[6]做"
        "接口服务，视频处理核心用FFmpeg，图像预处理用OpenCV，场景识别和语音识别这些AI功能用"
        "TensorFlow Lite在手机上跑。另外还会用到Git管代码、Postman测接口、Figma画界面、JProfiler"
        "做性能分析这些辅助工具。"
    )
    p14.runs[0].text = second_para
    p14.runs[0].font.size = Pt(12)

    refs = [
        '[1]\t中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R]. 北京: CNNIC, 2024.',
        '[2]\t黄清, 彭天强, 李玲, 等. 基于卷积神经网络的多模态视频场景分割优化算法[J]. 计算机应用研究, 2021, 38(10): 3134-3138.',
        '[3]\t王晓红, 包圆圆, 吕强. 移动短视频的发展现状及趋势观察[J]. 中国编辑, 2015, (3): 7-12.',
        '[4]\tFFmpeg Development Team. FFmpeg Documentation[EB/OL]. https://ffmpeg.org/documentation.html, 2024.',
        '[5]\tGoogle LLC. Flutter Documentation[EB/OL]. https://docs.flutter.dev/, 2024.',
        '[6]\t黑马程序员. Spring Boot企业级开发教程[M]. 北京: 人民邮电出版社, 2019.',
    ]
    for i, ref_text in enumerate(refs):
        p_idx = 16 + i
        p = doc.paragraphs[p_idx]
        for run in p.runs:
            run.text = ''
        p.runs[0].text = ref_text
        p.runs[0].font.size = Pt(12)

    doc.save('任务书.docx')
    print('任务书.docx 生成成功！')


def generate_opening_report():
    doc = Document('开题报告模板.docx')

    p16 = doc.paragraphs[16]
    for run in p16.runs:
        if run.text == ' 青':
            run.text = ' 移'
        elif run.text == '年':
            run.text = '动端短视频智能剪辑APP'
        elif run.text == '文化视频':
            run.text = ''
        elif run.text == '交流平台的设计与实现':
            run.text = '的开发与优化'

    p32 = doc.paragraphs[32]
    for run in p32.runs:
        if '2025' in run.text:
            run.text = run.text.replace('2025', '2026')

    table = doc.tables[0]

    for cell in [table.rows[0].cells[0], table.rows[0].cells[1]]:
        for p in cell.paragraphs:
            for run in p.runs:
                t = run.text
                if '青年文化视频交流平台的设计与实现' in t:
                    run.text = t.replace('青年文化视频交流平台的设计与实现', '移动端短视频智能剪辑APP的开发与优化')

    # Row 1: 选题目的和意义 (模板约413字)
    purpose_parts = [
        "一、选题的目的和意义（选题目的、理论价值依据、实践价值研究意义）",
        "",
        "这几年短视频发展得特别快，CNNIC发布的第53次统计报告显示，到2023年底国内短视频用户已经超过10亿了。很多人想自己剪视频，但市面上的工具不太好用，专业软件学不会，简单工具又不够智能，场景识别和片段筛选这些还得手动操作。所以本课题打算做一个手机上用的短视频智能剪辑APP，帮普通用户降低剪辑门槛。",
        "理论方面，这个课题想探索怎么把CNN+LSTM这种深度学习模型通过TensorFlow Lite部署到手机上来做场景识别，对移动端AI应用的开发能提供一些参考。实践方面，APP把视频编辑、场景识别、片段筛选、配乐推荐、字幕生成整合在一起，前端用Flutter做跨平台，后端用Spring Boot写接口，视频处理用FFmpeg，兼顾了开发效率和运行性能。",
        "希望做出来能让用户用着方便一些，有一定实用价值和应用前景。"
    ]
    for cell in [table.rows[1].cells[0], table.rows[1].cells[1]]:
        set_cell_content(cell, purpose_parts)

    # Row 2: 国内外研究状况综述 (模板约602字)
    review_parts = [
        "二、国内外研究状况综述（文献调研综述，对已有相关代表性研究成果的综合介绍与评价）",
        "",
        "国内研究",
        "国内这块比较有代表性的就是字节跳动的剪映了，功能确实挺丰富的，滤镜特效模板什么的都有，也做了智能字幕这些AI功能。不过它很多智能化的处理是在云端完成的，没网的时候就用不了。快手的快影也是用得比较多的剪辑工具，但它主打的是快捷简单，功能上偏基础一些。学术方面，王晓红等在《中国编辑》上写过移动短视频的发展现状分析，把行业整体情况梳理了一遍。技术方面，FFmpeg是做视频编解码绕不开的开源工具，国内很多视频APP底层都在用它。CNN和LSTM的组合在视频场景识别领域也有不少研究成果，黄清等人就研究过基于卷积神经网络的视频场景分割算法。",
        "国外研究",
        "国外的话，Adobe有个Premiere Rush是专门给手机用户做的轻量版剪辑工具，操作比电脑版简单不少。苹果自带的iMovie也是比较经典的移动端剪辑软件，界面比较好上手。技术上，Google的Flutter框架2018年出来以后发展很快，现在做跨平台APP用它的人越来越多了。TensorFlow Lite让深度学习模型可以在手机上跑起来，这对在移动端部署CNN+LSTM这类模型来说挺关键的。OpenCV在图像预处理方面也用得很广。",
        "总的来说，相关的工具和技术都不少了，但是怎么在手机上用比较简单的方案把智能剪辑做得好用一点，还是有空间可以做的。"
    ]
    for cell in [table.rows[2].cells[0], table.rows[2].cells[1]]:
        set_cell_content(cell, review_parts)

    # Row 3: 研究内容、方法、技术路线 (模板约769字)
    method_parts = [
        "三、研究的主要内容、研究方法、技术路线（包括基本思路、框架、主要研究方式或技术路线、方法、拟解决的关键问题、实验方案及可行性分析等）",
        "",
        "主要内容",
        "本课题要做的是一个手机上用的短视频智能剪辑APP，目标是让普通用户也能比较方便地剪辑视频。具体的研究内容有这几块：首先是基础编辑功能，包括视频导入（从相册选或者直接拍）、片段的裁剪和拼接、加各种转场效果、背景音乐添加和人声分离；然后是智能辅助的部分，用CNN+LSTM深度学习模型来自动识别视频中的不同场景（风景、人物、动态等），根据每帧画面的清晰度挑出质量比较好的片段，根据场景识别的结果来推荐合适的配乐；还有字幕自动生成功能，用语音识别把说话的内容转成文字；性能优化方面要调FFmpeg参数让导出更快，压缩图片减少内存占用，优化界面渲染减少卡顿；最后还有模板功能（节日、vlog、教程等）、画面比例切换和多格式导出。",
        "研究方法",
        "研究方法上，会先查阅国内外的论文和技术文档，搞清楚这个方向目前做到什么程度了。然后分析几个主流产品比如剪映和快影，看它们功能上有什么优缺点，了解用户实际的需求是什么。在这个基础上做需求分析和架构设计，然后进入编码实现阶段，开发完了反复测试和优化，确保功能好用、性能过关。",
        "技术路线",
        "系统整体是前后端分离的。前端用Flutter框架来写，它是用Dart语言的，好处是一套代码Android和iOS两个平台都能跑，不用写两遍。后端用Spring Boot搭RESTful风格的API接口，负责用户管理、模板数据、音乐资源这些后台功能，数据库用MySQL来存。视频处理的核心部分用FFmpeg来做编解码、裁剪拼接和格式转换，图像预处理用OpenCV，场景识别用CNN+LSTM模型通过TensorFlow Lite部署到手机上运行，语音识别也是同样用TensorFlow Lite来跑轻量化模型。",
        "功能框架"
    ]
    for cell in [table.rows[3].cells[0], table.rows[3].cells[1]]:
        set_cell_content(cell, method_parts)

    # Row 4: 研究特色与创新点 (模板约264字)
    innovation_parts = [
        "研究特色与创新点（概况和总结研究成果差异性，创新之处）",
        "",
        "场景识别用了CNN+LSTM深度学习模型，通过TensorFlow Lite部署到手机端运行，能自动识别风景、人物、动态场景等不同类型的画面，比单纯靠画面变化检测识别得更准。",
        "片段筛选是看视频每一帧的清晰度，自动帮用户挑出画质好的部分，这样用户拍了一堆素材不用一个一个自己去看了。",
        "用Flutter做前端，一套代码两个平台都能用，然后配合FFmpeg参数调优和图像压缩来保证手机上跑起来的性能还行，算是在跨平台和性能之间找了个平衡点。"
    ]
    for cell in [table.rows[4].cells[0], table.rows[4].cells[1]]:
        set_cell_content(cell, innovation_parts)

    # Row 5: 参考文献
    refs_parts = [
        "五、主要参考文献",
        "",
        "[1] 中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R]. 北京: CNNIC, 2024.",
        "[2] 王晓红, 包圆圆, 吕强. 移动短视频的发展现状及趋势观察[J]. 中国编辑, 2015, (3): 7-12.",
        "[3] Bradski G, Kaehler A. Learning OpenCV: Computer Vision with the OpenCV Library[M]. Sebastopol: O'Reilly Media, 2008.",
        "[4] FFmpeg Development Team. FFmpeg Documentation[EB/OL]. https://ffmpeg.org/documentation.html, 2024.",
        "[5] Google LLC. Flutter Documentation[EB/OL]. https://docs.flutter.dev/, 2024.",
        "[6] 黑马程序员. Spring Boot企业级开发教程[M]. 北京: 人民邮电出版社, 2019.",
        "[7] Google LLC. TensorFlow Lite Guide[EB/OL]. https://www.tensorflow.org/lite/guide, 2024.",
        "[8] OpenCV Team. OpenCV Documentation[EB/OL]. https://docs.opencv.org/, 2024.",
        "[9] 汪云飞, 沈永林, 陈晓茜. 从企业级开发到云原生微服务: Spring Boot实战[M]. 北京: 电子工业出版社, 2020.",
        "[10] Marcelo Luis Barbosa dos Santos. The \"so-called\" UGC: an updated definition of user-generated content in the age of social media[J]. Online Information Review, 2022, 46(1): 95-113.",
        "[11] 黄清, 彭天强, 李玲, 等. 基于卷积神经网络的多模态视频场景分割优化算法[J]. 计算机应用研究, 2021, 38(10): 3134-3138.",
        "[12] Dart Team. Dart Programming Language[EB/OL]. https://dart.dev/, 2024."
    ]
    for cell in [table.rows[5].cells[0], table.rows[5].cells[1]]:
        set_cell_content(cell, refs_parts, font_size=Pt(11))

    # Row 6: 进度计划
    plan_parts = [
        "六、完成措施及进度计划安排",
        "",
        "2025年11月20日 — 2025年12月07日 确定毕设选题",
        "2025年12月08日 — 2025年12月29日 完成任务书和开题报告",
        "2025年12月30日 — 2026年03月04日 进行专业实习，完成实习报告",
        "2026年03月05日 — 2026年03月19日 搜集、查阅资料，完成外文文献翻译",
        "2026年03月20日 — 2026年04月04日 分析系统，确定设计思路，设计出详细方案",
        "2026年04月05日 — 2026年05月20日 按照设计方案实施，完成系统设计，撰写论文",
        "2026年05月21日 — 2026年06月12日 准备论文答辩"
    ]
    for cell in [table.rows[6].cells[0], table.rows[6].cells[1]]:
        set_cell_content(cell, plan_parts)

    # Row 7: 指导教师意见
    teacher_parts = [
        "指导教师对开题报告的意见",
        "", "", "", "", "", "", "", "",
        "                                                       指导教师（签字）：",
        "2026年3月3日"
    ]
    for cell in [table.rows[7].cells[0], table.rows[7].cells[1]]:
        set_cell_content(cell, teacher_parts)

    # Row 8: 学院审核
    review_opinion_parts = [
        "学院审核意见（在相应处打勾）",
        "",
        "通过；         2、未通过",
        "",
        "专业负责人（签字）：",
        "2026年3月4日"
    ]
    for cell in [table.rows[8].cells[0], table.rows[8].cells[1]]:
        set_cell_content(cell, review_opinion_parts)

    # 去除所有行的固定高度，让行高自动适应内容
    for row in table.rows:
        tr = row._element
        trPr = tr.find(qn('w:trPr'))
        if trPr is not None:
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                trPr.remove(trHeight)

    doc.save('开题报告.docx')
    print('开题报告.docx 生成成功！')


if __name__ == '__main__':
    generate_task_book()
    generate_opening_report()
