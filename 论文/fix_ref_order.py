"""
修复参考文献引用顺序：按正文中首次出现的顺序重新编号为 1,2,3,...
"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import re
import copy
import shutil

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"
BAK = DOC_PATH + ".bak_before_reorder"
shutil.copy2(DOC_PATH, BAK)
print(f"备份: {BAK}")

doc = Document(DOC_PATH)

ref_section_start = None
ref_paras = {}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '参考文献' in text and len(text) < 20 and i > 100:
        ref_section_start = i
        print(f"参考文献标题: P{i}")
        continue
    if ref_section_start and i > ref_section_start:
        m = re.match(r'^\[(\d+)\]', text)
        if m:
            num = int(m.group(1))
            ref_paras[num] = i
        elif text and not text.startswith('[') and len(ref_paras) > 0:
            break

print(f"找到 {len(ref_paras)} 条参考文献: {sorted(ref_paras.keys())}")

cite_order = []
seen = set()
body_end = ref_section_start if ref_section_start else len(doc.paragraphs)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

for i, p in enumerate(doc.paragraphs):
    if i >= body_end:
        break

    for elem in p._element.iter():
        if elem.tag == qn('w:instrText') or (isinstance(elem.tag, str) and 'instrText' in elem.tag):
            instr = elem.text or ''
            m = re.search(r'REF\s+_Ref_(\d+)', instr)
            if m:
                num = int(m.group(1))
                if num not in seen:
                    cite_order.append(num)
                    seen.add(num)

    text = p.text
    if text:
        for m in re.finditer(r'\[(\d+)\]', text):
            num = int(m.group(1))
            if num not in seen and num in ref_paras:
                cite_order.append(num)
                seen.add(num)

for num in sorted(ref_paras.keys()):
    if num not in seen:
        cite_order.append(num)
        seen.add(num)

print(f"引用顺序（前20）: {cite_order[:20]}...")

old_to_new = {}
for new_num, old_num in enumerate(cite_order, 1):
    old_to_new[old_num] = new_num
print(f"映射（old→new）: { {k: old_to_new[k] for k in sorted(old_to_new.keys())[:10]} }...")

for elem in doc.element.iter():
    if elem.tag == qn('w:instrText') or (isinstance(elem.tag, str) and 'instrText' in elem.tag):
        instr = elem.text or ''
        m = re.search(r'(REF\s+_Ref_)(\d+)', instr)
        if m:
            old_num = int(m.group(2))
            if old_num in old_to_new:
                new_num = old_to_new[old_num]
                elem.text = instr.replace(f'_Ref_{old_num}', f'_Ref_NEW_{new_num}')

for elem in doc.element.iter():
    if elem.tag == qn('w:instrText') or (isinstance(elem.tag, str) and 'instrText' in elem.tag):
        if elem.text:
            elem.text = elem.text.replace('_Ref_NEW_', '_Ref_')

updated_runs = 0
for elem in doc.element.iter():
    tag_local = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ''
    if tag_local == 't' and elem.text:
        text = elem.text
        parent = elem.getparent()
        if parent is not None:
            rPr = parent.find(qn('w:rPr'))
            if rPr is not None:
                va = rPr.find(qn('w:vertAlign'))
                if va is not None and va.get(qn('w:val')) == 'superscript':
                    new_text = text
                    for m in re.finditer(r'\[(\d+)\]', text):
                        old_num = int(m.group(1))
                        if old_num in old_to_new:
                            new_text = new_text.replace(f'[{old_num}]', f'[NEW{old_to_new[old_num]}]')
                    if new_text != text:
                        new_text = new_text.replace('[NEW', '[')
                        elem.text = new_text
                        updated_runs += 1

print(f"更新了 {updated_runs} 个上标引用文本")

for bm_start in doc.element.iter(qn('w:bookmarkStart')):
    name = bm_start.get(qn('w:name'))
    if name and name.startswith('_Ref_'):
        m = re.match(r'_Ref_(\d+)', name)
        if m:
            old_num = int(m.group(1))
            if old_num in old_to_new:
                bm_start.set(qn('w:name'), f'_Ref_TEMP_{old_to_new[old_num]}')

for bm_start in doc.element.iter(qn('w:bookmarkStart')):
    name = bm_start.get(qn('w:name'))
    if name and name.startswith('_Ref_TEMP_'):
        bm_start.set(qn('w:name'), name.replace('_Ref_TEMP_', '_Ref_'))

print("书签已更新")

ref_contents = {}
for old_num, para_idx in ref_paras.items():
    p = doc.paragraphs[para_idx]
    old_text = p.text.strip()
    content = re.sub(r'^\[\d+\]\s*', '', old_text)
    ref_contents[old_num] = content

for old_num, para_idx in ref_paras.items():
    new_num = old_to_new[old_num]
    p = doc.paragraphs[para_idx]
    old_text = p.text.strip()
    new_prefix = f'[{new_num}] '
    content = ref_contents[old_num]
    
    for run in p.runs:
        old_run_text = run.text
        m = re.match(r'^\[(\d+)\]\s*', old_run_text)
        if m:
            run.text = new_prefix + old_run_text[m.end():]
            break
        m2 = re.match(r'^\[(\d+)\]', old_run_text)
        if m2:
            run.text = f'[{new_num}]' + old_run_text[m2.end():]
            break

print("参考文献编号已更新")

doc.save(DOC_PATH)
print(f"文档已保存: {DOC_PATH}")
print("\n== 完成 ==")
print(f"参考文献重排映射:")
for old_num in cite_order:
    new_num = old_to_new[old_num]
    content_preview = ref_contents.get(old_num, '???')[:50]
    if old_num != new_num:
        print(f"  [{old_num}] → [{new_num}]: {content_preview}...")
    else:
        print(f"  [{old_num}] = [{new_num}]: {content_preview}...")
