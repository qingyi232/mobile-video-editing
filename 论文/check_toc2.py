# -*- coding: utf-8 -*-
import docx
from lxml import etree

doc = docx.Document(r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v4.docx')

lines = []

# Check paragraphs around where TOC should be
for i in range(120, 130):
    p = doc.paragraphs[i]
    t = p.text.strip()
    style = p.style.name if p.style else 'None'
    
    # Check for field codes (TOC)
    fld_chars = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
    instr_texts = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText')
    
    fld_info = ''
    if fld_chars:
        fld_types = [fc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', '?') for fc in fld_chars]
        fld_info = f' FIELD_CHARS={fld_types}'
    if instr_texts:
        instr = ''.join(it.text or '' for it in instr_texts)
        fld_info += f' INSTR="{instr.strip()}"'
    
    lines.append(f'[{i}] style={style} | text="{t[:60]}"{fld_info}')

# Also search entire document for TOC field
lines.append('\n=== Searching for TOC field across document ===')
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for i, p in enumerate(doc.paragraphs):
    instr_texts = p._element.findall('.//w:instrText', ns)
    if instr_texts:
        instr = ''.join(it.text or '' for it in instr_texts)
        if 'TOC' in instr.upper() or '目录' in instr:
            lines.append(f'[{i}] TOC FIELD: {instr.strip()}')

# Search for "目录" text
lines.append('\n=== Searching for 目录 text ===')
for i, p in enumerate(doc.paragraphs):
    if '目录' in p.text:
        lines.append(f'[{i}] {p.text.strip()[:60]}')

with open(r'F:\26毕设2\移动端短视频智能剪辑app\论文\toc_field_info.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print('Done')
