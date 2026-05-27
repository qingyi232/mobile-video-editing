# -*- coding: utf-8 -*-
import docx
from docx.oxml.ns import qn

doc_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v3.docx'
out_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v4.docx'

doc = docx.Document(doc_path)

para_228 = (
    "在技术实现层面，视频处理功能的底层调用统一封装在FFmpegUtil工具类中，"
    "该类通过Java的ProcessBuilder API以子进程方式调用FFmpeg命令行工具。"
    "所有FFmpeg操作方法遵循统一的设计模式：首先根据功能需求构建FFmpeg命令参数列表，"
    "然后通过executeCommand方法创建子进程执行命令，该方法将标准输出和错误输出合并读取以防止管道缓冲区满导致进程阻塞，"
    "最终通过检查进程退出码判断操作是否成功。对于需要获取命令输出内容的场景（如通过ffprobe获取视频时长、分辨率等元数据信息），"
    "则使用executeCommandWithOutput方法将子进程的完整输出捕获到StringBuilder中，"
    "再通过正则表达式从JSON格式的输出中解析出所需的字段值。"
    "这种封装设计将FFmpeg的命令行复杂性屏蔽在工具类内部，使上层的VideoProjectService只需关注业务逻辑。"
)

para_229 = (
    "在视频文件处理过程中，系统采用了多项关键技术策略来保障处理的可靠性和效率。"
    "所有视频编解码操作统一使用libx264视频编码器和AAC音频编码器，并配置fast预设参数以平衡编码质量与处理速度。"
    "视频拼接功能采用FFmpeg的concat协议实现，通过创建临时文本文件记录待拼接视频的路径列表，"
    "处理完成后自动清理临时文件，避免了直接在命令行中传递大量路径参数带来的长度限制问题。"
    "转场特效功能利用FFmpeg的xfade视频滤镜，通过filter_complex参数将两个输入视频流组合，"
    "支持fade、slideleft、smoothup、circleopen和wipeleft五种转场效果。"
    "背景音乐添加功能采用双重降级策略：首先尝试使用amix滤镜混合原始视频音轨与背景音乐，"
    "若原视频无音轨则自动降级为仅叠加背景音乐轨道，确保不同来源视频的兼容处理。"
    "人声分离功能则通过组合使用highpass和lowpass频率滤波器，"
    "以300Hz和3000Hz为截止频率分别提取人声频段和背景音频段。"
)

# Get reference formatting from paragraph 227
ref_para = doc.paragraphs[227]
ref_style = ref_para.style
ref_run = ref_para.runs[0] if ref_para.runs else None

print(f"Reference style: {ref_style.name}")
if ref_run:
    print(f"Reference font: name={ref_run.font.name}, size={ref_run.font.size}")

def set_paragraph_text(para, text, ref_run_obj):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        run = para.add_run(text)
    
    run = para.runs[0]
    if ref_run_obj:
        if ref_run_obj.font.name:
            run.font.name = ref_run_obj.font.name
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = docx.oxml.OxmlElement('w:rFonts')
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), ref_run_obj.font.name)
        if ref_run_obj.font.size:
            run.font.size = ref_run_obj.font.size

# Fill paragraph 228
p228 = doc.paragraphs[228]
print(f"P228 before: '{p228.text.strip()[:30]}' (len={len(p228.text.strip())})")
set_paragraph_text(p228, para_228, ref_run)
print(f"P228 after: len={len(p228.text.strip())}")

# Fill paragraph 229
p229 = doc.paragraphs[229]
print(f"P229 before: '{p229.text.strip()[:30]}' (len={len(p229.text.strip())})")
set_paragraph_text(p229, para_229, ref_run)
print(f"P229 after: len={len(p229.text.strip())}")

# Remove empty paragraph 230 if it exists and is empty
p230 = doc.paragraphs[230]
if not p230.text.strip():
    p230._element.getparent().remove(p230._element)
    print("Removed empty P230")

doc.save(out_path)
print(f"\nSaved to: {out_path}")

# Verify
doc2 = docx.Document(out_path)
print("\n=== Verification (226-232) ===")
for i in range(226, 233):
    p = doc2.paragraphs[i]
    t = p.text.strip()
    print(f"[{i}] ({len(t)} chars) {t[:80]}")
