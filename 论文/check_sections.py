# -*- coding: utf-8 -*-
import docx
import re

lines = []

for label, doc_path in [
    ('ORIG', r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改.docx'),
    ('V2', r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v2.docx'),
    ('V3', r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v3.docx'),
]:
    doc = docx.Document(doc_path)
    lines.append('=== {} (total: {} paras) ==='.format(label, len(doc.paragraphs)))
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if re.match(r'^\d+(\.\d+)*\s', t) and len(t) < 50:
            lines.append('  [{}] {}'.format(i, t))
    lines.append('')

with open(r'F:\26毕设2\移动端短视频智能剪辑app\论文\sections_compare.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('Done')
