# -*- coding: utf-8 -*-
import docx

lines = []

for label, doc_path in [
    ('ORIG', r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改.docx'),
    ('V3', r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v3.docx'),
]:
    doc = docx.Document(doc_path)
    lines.append('=== {} - Section 4.2.2 Video Processing ==='.format(label))
    
    start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if '4.2.2' in t and '视频处理' in t:
            start = i
        if start and ('4.2.3' in t or '音乐推荐' in t) and i > start:
            end = i
            break
    
    if start:
        for j in range(start, end):
            p = doc.paragraphs[j]
            t = p.text.strip()
            lines.append('[{}] ({} chars) {}'.format(j, len(t), t[:150]))
    lines.append('')

with open(r'F:\26毕设2\移动端短视频智能剪辑app\论文\vidproc_compare.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))
print('Done')
