"""补充剩余6条未引用的参考文献"""
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"
doc = Document(DOC_PATH)

BODY_START = 125
BODY_END = 480

def find_para_in_body(doc, search_text, start=BODY_START, end=BODY_END):
    for i, p in enumerate(doc.paragraphs):
        if i < start or i > end:
            continue
        if search_text in p.text:
            return i
    return None

def add_sup_ref(doc, pidx, ref_nums):
    if pidx is None:
        return False
    para = doc.paragraphs[pidx]
    ref_text = '[' + ','.join(str(n) for n in ref_nums) + ']'
    run = para.add_run(ref_text)
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="SimSun"/>')
    rPr.insert(0, rFonts)
    vertAlign = parse_xml(f'<w:vertAlign {nsdecls("w")} w:val="superscript"/>')
    rPr.append(vertAlign)
    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="18"/>')
    szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="18"/>')
    rPr.append(sz)
    rPr.append(szCs)
    return True

# [6] Li comprehensive survey on video summarization
# [7] Altundogan multi-objective video summarization
# [10] 门飞 智能音视频剪辑
# [12] Teng smooth video summarization
# [16] Radford speech recognition (百度ASR)
# [17] TensorFlow Lite

# First find available targets
print("Searching for remaining ref targets...")
targets = [
    # [6] Li - 对视频摘要进行全面综述
    ('视频摘要技术', [6], '国外研究'),
    # [7] Altundogan - 多目标视频摘要
    ('视频编辑器的功能虽然十分强大', [7], '专业视频工具'),
    # [10] 门飞 - 智能音视频剪辑自动化
    ('人声分离功能', [10], '人声分离'),
    # [12] Teng - 帧填充平滑视频摘要
    ('实现了一键智能剪辑', [12], '智能剪辑'),
    # [16] Radford - 语音识别
    ('语音识别服务类', [16], '百度ASR'),
    # [17] TensorFlow Lite
    ('CNN+LSTM模型', [17], '模型部署'),
]

count = 0
for search_text, ref_nums, desc in targets:
    pidx = find_para_in_body(doc, search_text)
    if pidx is not None:
        success = add_sup_ref(doc, pidx, ref_nums)
        if success:
            count += 1
            safe = doc.paragraphs[pidx].text[:50].encode('gbk', errors='replace').decode('gbk')
            print(f"  OK [{ref_nums}] -> P{pidx} ({desc}): {safe}...")
    else:
        print(f"  MISS [{ref_nums}]: '{search_text}' ({desc})")

# Try alternate searches for missed refs
if count < len(targets):
    alternates = [
        ('Li等人', [6], 'Li survey alt1'),
        ('全面的综述', [6], 'Li survey alt2'),
        ('Adobe Premiere Pro', [7], 'Altundogan alt'),
        ('人声和背景音分别提取', [10], '门飞 alt'),
        ('帧填充', [12], 'Teng alt'),
        ('将视频切割为多个音频片段', [16], 'Radford alt'),
        ('均匀采样', [17], 'TFLite alt'),
    ]
    for search_text, ref_nums, desc in alternates:
        if ref_nums[0] not in [r for _, refs, _ in targets if find_para_in_body(doc, _) is not None for r in refs]:
            pidx = find_para_in_body(doc, search_text)
            if pidx is not None:
                success = add_sup_ref(doc, pidx, ref_nums)
                if success:
                    count += 1
                    safe = doc.paragraphs[pidx].text[:50].encode('gbk', errors='replace').decode('gbk')
                    print(f"  OK-ALT [{ref_nums}] -> P{pidx} ({desc}): {safe}...")

print(f"\nTotal supplemental: {count}")

doc.save(DOC_PATH)
print("Saved.")
