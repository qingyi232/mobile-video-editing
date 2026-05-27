import docx
import os

doc_path = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20.docx"
doc = docx.Document(doc_path)

output_lines = []

output_lines.append("=== 文档段落结构 ===\n")
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else "None"
    text = para.text.strip()
    if text:
        has_image = False
        for run in para.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                has_image = True
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
                has_image = True
        
        if 'Heading' in style_name or '标题' in style_name or text.startswith('第') or '图' in text[:6] or '表' in text[:6] or '公式' in text[:6]:
            output_lines.append(f"P{i} [{style_name}] {'[IMG]' if has_image else ''} {text[:120]}")
        elif len(text) < 80 and ('图' in text or '表' in text):
            output_lines.append(f"P{i} [{style_name}] {'[IMG]' if has_image else ''} {text[:120]}")

output_lines.append("\n\n=== 所有图片相关段落 ===\n")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    has_drawing = False
    for run in para.runs:
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if drawings:
            has_drawing = True
        inline_drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if inline_drawings:
            has_drawing = True
    
    if has_drawing or ('图' in text[:10] and ('.' in text[:10] or '：' in text[:10] or ' ' in text[:15])):
        output_lines.append(f"P{i} [has_drawing={has_drawing}] {text[:150]}")

output_lines.append("\n\n=== 带图片的段落详情 ===\n")
from docx.oxml.ns import qn
for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall('.//' + qn('wp:inline'))
    anchors = para._element.findall('.//' + qn('wp:anchor'))
    if drawings or anchors:
        text = para.text.strip()
        output_lines.append(f"P{i} inline={len(drawings)} anchor={len(anchors)} text='{text[:100]}'")
        prev_text = doc.paragraphs[i-1].text.strip() if i > 0 else ""
        next_text = doc.paragraphs[i+1].text.strip() if i < len(doc.paragraphs)-1 else ""
        output_lines.append(f"  prev: {prev_text[:100]}")
        output_lines.append(f"  next: {next_text[:100]}")

output_lines.append("\n\n=== 第3章和第4章完整内容 ===\n")
in_target = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style_name = para.style.name if para.style else "None"
    
    if '第3章' in text or '第三章' in text or (text.startswith('3') and 'Heading' in style_name):
        in_target = True
    if '第6章' in text or '第六章' in text or (text.startswith('6') and 'Heading' in style_name):
        in_target = False
    
    if in_target:
        has_drawing = bool(para._element.findall('.//' + qn('wp:inline')) or para._element.findall('.//' + qn('wp:anchor')))
        marker = "[IMG]" if has_drawing else ""
        output_lines.append(f"P{i} [{style_name}] {marker} {text[:200]}")

result = "\n".join(output_lines)
with open(r"F:\26毕设2\移动端短视频智能剪辑app\论文\doc_structure.txt", "w", encoding="utf-8") as f:
    f.write(result)

print(f"Done. {len(output_lines)} lines written.")
