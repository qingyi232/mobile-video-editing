"""
将生成的图片替换到论文文档中，并插入公式。
不修改其他任何内容和格式。
"""
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20.docx"
FIG_DIR = r"F:\26毕设2\移动端短视频智能剪辑app\论文\figures"
OUTPUT_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"

shutil.copy2(DOC_PATH, DOC_PATH + '.bak_before_fig2')

doc = Document(DOC_PATH)
body = doc.element.body

def find_para_idx(doc, search_text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search_text in p.text:
            return i
    return None

def find_image_para_before_caption(doc, caption_prefix):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(caption_prefix) and i > 0:
            prev = doc.paragraphs[i - 1]
            if prev._element.findall('.//' + qn('wp:inline')):
                return i - 1
    return None

# ============================================================
# Step 1: 替换现有图片
# ============================================================
replacements = {
    '图 3.1': ('fig31_architecture.png', Cm(15)),
    '图 3.2': ('fig32_modules.png', Cm(16)),
    '图 3.3': ('fig33_er.png', Cm(15)),
    '图 4.27': ('fig427_workflow.png', Cm(14)),
}

for cap_key, (img_file, img_width) in replacements.items():
    img_path = os.path.join(FIG_DIR, img_file)
    pidx = find_image_para_before_caption(doc, cap_key)
    if pidx is not None:
        para = doc.paragraphs[pidx]
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        run = para.add_run()
        run.add_picture(img_path, width=img_width)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"OK 替换 {cap_key} (P{pidx})")
    else:
        print(f"WARN 未找到 {cap_key}")

# ============================================================
# Step 2: 插入后端服务流程图
# 
# 策略：先用 doc.add_paragraph 创建段落（这样 part 引用正确），
# 然后把元素移动到正确位置。
# ============================================================
flow_insertions = [
    ('系统通过JwtAuthenticationFilter过滤器来实现对API请求的统一认证拦截', 'fig_flow_auth.png', '图 4.28  用户认证模块流程图', Cm(13)),
    ('背景音乐添加功能采用双重降级策略', 'fig_flow_video.png', '图 4.29  视频处理模块流程图', Cm(14)),
    ('识别结果以JSON格式返回给前端', 'fig_flow_scene.png', '图 4.30  场景识别模块流程图', Cm(13)),
    ('原始视频文件不受影响', 'fig_flow_subtitle.png', '图 4.31  字幕生成模块流程图', Cm(13)),
]

for search_text, img_file, caption_text, width in reversed(flow_insertions):
    pidx = find_para_idx(doc, search_text)
    if pidx is None:
        print(f"WARN 未找到 '{search_text[:30]}...'")
        continue

    ref_elem = doc.paragraphs[pidx]._element

    # Create caption paragraph at end of doc, then move
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption_text)
    cap_run.font.name = 'SimSun'
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    cap_run.font.size = Pt(10.5)
    # Try to use Caption style
    try:
        cap_para.style = doc.styles['Caption']
    except:
        pass
    # Move after ref
    cap_elem = cap_para._element
    body.remove(cap_elem)
    ref_elem.addnext(cap_elem)

    # Create image paragraph at end of doc, then move
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_run = img_para.add_run()
    img_path = os.path.join(FIG_DIR, img_file)
    img_run.add_picture(img_path, width=width)
    # Move between ref and caption
    img_elem = img_para._element
    body.remove(img_elem)
    ref_elem.addnext(img_elem)

    print(f"OK 插入 {caption_text} (after P{pidx})")

# ============================================================
# Step 3: 插入公式
# ============================================================
def insert_formula_after(doc, after_text, formula_text, formula_num=''):
    pidx = find_para_idx(doc, after_text)
    if pidx is None:
        print(f"WARN 公式位置未找到: '{after_text[:30]}...'")
        return
    
    ref_elem = doc.paragraphs[pidx]._element
    
    f_para = doc.add_paragraph()
    
    # Set paragraph format: spacing before/after 6pt, single line spacing
    pf = f_para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(12)
    pf.first_line_indent = Cm(0)
    
    # Add formula text
    f_run = f_para.add_run(formula_text)
    f_run.font.name = 'Times New Roman'
    f_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    f_run.font.size = Pt(12)
    
    if formula_num:
        # Add tab stop at right margin for formula number
        tab_run = f_para.add_run('\t' + formula_num)
        tab_run.font.name = 'Times New Roman'
        tab_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        tab_run.font.size = Pt(12)
        
        # Add right tab stop
        pPr = f_para._element.pPr
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            f_para._element.insert(0, pPr)
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="8300"/></w:tabs>')
        pPr.append(tabs)
    
    # Move to correct position
    f_elem = f_para._element
    body.remove(f_elem)
    ref_elem.addnext(f_elem)

# Scene recognition CNN formulas (insert in reverse order so they appear correct)
formulas_cnn = [
    ('模拟卷积神经网络对单帧图像空间特征的编码过程', [
        ('其中，L(I)为Laplacian算子对图像I的卷积结果，Var(·)表示方差运算', ''),
        ('Clarity(I) = Var(L(I)) = Var(∇²I)', '（4-1）'),
    ]),
    ('用于捕捉视频内容的渐变或突变模式', [
        ('Trend = mean(Clarity_后半段) - mean(Clarity_前半段)', '（4-4）'),
        ('Motion_avg = (1/N) · Σ M(t),  Motion_max = max{M(t)}', '（4-3）'),
        ('M(t) = |Clarity(t) - Clarity(t-1)|,  t = 2, 3, ..., N', '（4-2）'),
    ]),
    ('将其限制在0.55至0.95的合理范围内', [
        ('Confidence = clip(Score_max / Σ Score_k, 0.55, 0.95)', '（4-6）'),
        ('Score_k = Σ(w_i · f_i),  k ∈ {风景, 人像, 运动, 静态}', '（4-5）'),
    ]),
    ('按照使用频率和评分等指标进行排序后推荐给用户', [
        ('其中，α + β + γ = 1，Match为场景类别匹配函数，Freq为使用频率，Rating为评分', ''),
        ('Relevance(m, s) = α · Match(category_m, scene_s) + β · Freq(m) + γ · Rating(m)', '（4-7）'),
    ]),
]

for after_text, formula_list in formulas_cnn:
    for ftxt, fnum in formula_list:
        insert_formula_after(doc, after_text, ftxt, fnum)
    print(f"OK 公式组 after '{after_text[:30]}...'")

# ============================================================
# Save
# ============================================================
doc.save(OUTPUT_PATH)
print(f"\n文档已保存到:\n{OUTPUT_PATH}")
