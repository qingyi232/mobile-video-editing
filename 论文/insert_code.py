"""
在后端服务流程图之后插入关键代码段。
操作的是修改版文档。
"""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"

doc = Document(DOC_PATH)
body = doc.element.body

def find_para_idx(doc, search_text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search_text in p.text:
            return i
    return None

def insert_code_block(doc, after_text, title, code_lines, desc_before=None):
    """Insert a titled code block after the specified text."""
    pidx = find_para_idx(doc, after_text)
    if pidx is None:
        print(f"WARN not found: '{after_text[:40]}...'")
        return

    ref_elem = doc.paragraphs[pidx]._element
    elements_to_insert = []

    if desc_before:
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.space_before = Pt(6)
        desc_para.paragraph_format.space_after = Pt(3)
        desc_para.paragraph_format.first_line_indent = Cm(0.74)
        r = desc_para.add_run(desc_before)
        r.font.name = 'SimSun'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.font.size = Pt(12)
        elements_to_insert.append(desc_para._element)

    for line in code_lines:
        code_para = doc.add_paragraph()
        code_para.paragraph_format.space_before = Pt(0)
        code_para.paragraph_format.space_after = Pt(0)
        code_para.paragraph_format.line_spacing = Pt(14)
        code_para.paragraph_format.first_line_indent = Cm(0)

        r = code_para.add_run(line)
        r.font.name = 'Courier New'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.font.size = Pt(9)

        # Light gray background for code
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr = code_para._element.pPr
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            code_para._element.insert(0, pPr)
        pPr.append(shd)

        elements_to_insert.append(code_para._element)

    # Move all elements from end of doc to correct position (in reverse order)
    for elem in elements_to_insert:
        body.remove(elem)

    for elem in reversed(elements_to_insert):
        ref_elem.addnext(elem)

    print(f"OK inserted code after '{after_text[:40]}...'")


# ============================================================
# 1. 用户认证关键代码 - 在图4.28标题之后
# ============================================================
insert_code_block(doc,
    '图 4.28  用户认证模块流程图',
    '用户认证关键代码',
    [
        '// UserService.java - 用户登录核心逻辑',
        'public LoginResponse login(LoginRequest request) {',
        '    User user = userRepository.findByUsername(',
        '            request.getUsername())',
        '        .orElseThrow(() -> new RuntimeException("用户不存在"));',
        '    if (!passwordEncoder.matches(',
        '            request.getPassword(), user.getPassword())) {',
        '        throw new RuntimeException("密码错误");',
        '    }',
        '    String token = jwtUtil.generateToken(',
        '            user.getId(), user.getUsername());',
        '    return new LoginResponse(token, user.getId(),',
        '            user.getUsername(), user.getNickname(),',
        '            user.getAvatar());',
        '}',
    ],
    desc_before='用户认证模块的关键代码如下所示，login方法通过BCrypt算法验证密码，验证通过后生成JWT令牌返回给客户端：'
)

# ============================================================
# 2. 视频处理关键代码 - 在图4.29标题之后
# ============================================================
insert_code_block(doc,
    '图 4.29  视频处理模块流程图',
    '视频裁剪关键代码',
    [
        '// VideoProjectService.java - 视频裁剪核心逻辑',
        'public String trimVideo(String inputPath,',
        '        double startTime, double endTime, Long userId) {',
        '    String absInput = toAbsolutePath(inputPath);',
        '    String dir = exportDir + "/" + userId;',
        '    new File(dir).mkdirs();',
        '    String outputPath = dir + "/trimmed_"',
        '        + UUID.randomUUID() + ".mp4";',
        '    boolean success = ffmpegUtil.trimVideo(',
        '        absInput, outputPath, startTime, endTime);',
        '    if (!success)',
        '        throw new RuntimeException("视频裁剪失败");',
        '    validateOutput(outputPath, "视频裁剪");',
        '    return toRelativePath(outputPath);',
        '}',
    ],
    desc_before='视频处理模块的关键代码如下所示。视频裁剪通过FFmpeg的-ss和-to参数实现精确截取，所有FFmpeg操作均通过ProcessBuilder以子进程方式执行：'
)

# ============================================================
# 3. 场景识别关键代码 - 在图4.30标题之后
# ============================================================
insert_code_block(doc,
    '图 4.30  场景识别模块流程图',
    '场景识别关键代码',
    [
        '// VideoProjectService.java - CNN+LSTM场景识别',
        'public Map<String, Object> recognizeScene(String videoPath) {',
        '    // 阶段一：CNN特征提取 - 均匀采样8帧',
        '    int sampleCount = 8;',
        '    List<Map<String, Object>> frameSamples',
        '        = analyzeClarity(videoPath, sampleCount);',
        '    double[] claritySeq = frameSamples.stream()',
        '        .mapToDouble(s -> ((Number) s.getOrDefault(',
        '            "clarity", 0)).doubleValue()).toArray();',
        '    double avgClarity = Arrays.stream(claritySeq)',
        '        .average().orElse(0);',
        '',
        '    // 阶段二：LSTM时序建模 - 计算帧间运动强度',
        '    double[] motionSeq = new double[claritySeq.length - 1];',
        '    for (int i = 1; i < claritySeq.length; i++)',
        '        motionSeq[i-1] = Math.abs(',
        '            claritySeq[i] - claritySeq[i-1]);',
        '',
        '    // 阶段三：全连接分类 - 加权评分',
        '    Map<String, Double> scores = new LinkedHashMap<>();',
        '    scores.put("nature", natureScore);   // 风景',
        '    scores.put("portrait", portraitScore);// 人像',
        '    scores.put("dynamic", dynamicScore);  // 运动',
        '    scores.put("calm", calmScore);        // 静态',
        '    double confidence = Math.max(0.55,',
        '        Math.min(0.95, topScore / totalScore));',
        '}',
    ],
    desc_before='场景识别模块采用CNN+LSTM轻量化方案，关键代码如下所示。系统通过Laplacian方差提取每帧清晰度特征，计算帧间运动强度，最后通过加权评分进行场景分类：'
)

# ============================================================
# 4. 字幕生成关键代码 - 在图4.31标题之后
# ============================================================
insert_code_block(doc,
    '图 4.31  字幕生成模块流程图',
    '字幕生成关键代码',
    [
        '// VideoProjectService.java - 字幕生成核心逻辑',
        'public Map<String, Object> generateSubtitle(',
        '        String videoPath) {',
        '    String absPath = toAbsolutePath(videoPath);',
        '    double duration = getDuration(absPath);',
        '    double windowSize = 15.0; // 每段15秒',
        '    StringBuilder fullText = new StringBuilder();',
        '    List<Map<String, Object>> segments',
        '        = new ArrayList<>();',
        '',
        '    for (double start = 0; start < duration;',
        '            start += windowSize) {',
        '        double end = Math.min(',
        '            start + windowSize, duration);',
        '        String text = baiduAsrService',
        '            .recognizeSegment(absPath, start, end);',
        '        if (text != null && !text.isEmpty()) {',
        '            fullText.append(text);',
        '            segments.add(Map.of(',
        '                "start", start, "end", end,',
        '                "text", text));',
        '        }',
        '    }',
        '    return Map.of("text", fullText.toString(),',
        '        "segments", segments);',
        '}',
    ],
    desc_before='字幕生成模块的关键代码如下所示。系统按15秒窗口切割视频，依次调用百度语音识别API进行语音转文字，最后按时间轴聚合生成完整字幕：'
)

doc.save(DOC_PATH)
print(f"\n关键代码已插入到文档: {DOC_PATH}")
