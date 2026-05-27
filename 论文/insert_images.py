# -*- coding: utf-8 -*-
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v2.docx'
img1_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\基础编辑标签页截图.png'
img2_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\智能辅助标签页截图.png'
out_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书修改_v3.docx'

doc = docx.Document(doc_path)

ref_caption = doc.paragraphs[281]
print(f"Reference caption style: {ref_caption.style.name}")
print(f"Reference caption text: {ref_caption.text.strip()}")

p270 = doc.paragraphs[270]
p273 = doc.paragraphs[273]
print(f"P270: {p270.text.strip()[:40]}")
print(f"P273: {p273.text.strip()[:40]}")


def insert_image_after(doc, target_para, img_path, width_cm=8):
    new_para = docx.oxml.OxmlElement('w:p')
    pPr = docx.oxml.OxmlElement('w:pPr')
    jc = docx.oxml.OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_para.append(pPr)
    target_para._element.addnext(new_para)

    new_paragraph = docx.text.paragraph.Paragraph(new_para, doc)
    run = new_paragraph.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return new_paragraph


def insert_caption_after(doc, target_para, caption_text, ref_caption_para):
    new_para = docx.oxml.OxmlElement('w:p')
    pPr = docx.oxml.OxmlElement('w:pPr')
    jc = docx.oxml.OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    if ref_caption_para.style and ref_caption_para.style.name != 'Default Paragraph Font':
        pStyle = docx.oxml.OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), ref_caption_para.style.style_id)
        pPr.append(pStyle)

    new_para.append(pPr)
    target_para._element.addnext(new_para)

    new_paragraph = docx.text.paragraph.Paragraph(new_para, doc)
    run = new_paragraph.add_run(caption_text)

    if ref_caption_para.runs:
        ref_run = ref_caption_para.runs[0]
        if ref_run.font.name:
            run.font.name = ref_run.font.name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), ref_run.font.name)
        if ref_run.font.size:
            run.font.size = ref_run.font.size
        run.font.bold = ref_run.font.bold

    return new_paragraph


print("\n--- Inserting images (reverse order to preserve indices) ---")

# Insert 智能辅助 after p273 (reverse: caption first, then image pushes above)
caption2 = insert_caption_after(doc, p273, "图4.5  智能辅助标签页", ref_caption)
img2_para = insert_image_after(doc, p273, img2_path, width_cm=8)
print("Inserted: 智能辅助标签页 → 图4.5 (after P273)")

# Insert 基础编辑 after p270
caption1 = insert_caption_after(doc, p270, "图4.4  基础编辑标签页", ref_caption)
img1_para = insert_image_after(doc, p270, img1_path, width_cm=8)
print("Inserted: 基础编辑标签页 → 图4.4 (after P270)")

print("\n--- Renaming captions (old→new) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '图4.4  视频编辑器界面':
        for run in p.runs:
            run.text = run.text.replace('图4.4', '图4.6')
        print(f"  [{i}] 图4.4 → 图4.6 (视频编辑器界面)")
    elif t == '图4.5  个人中心界面':
        for run in p.runs:
            run.text = run.text.replace('图4.5', '图4.7')
        print(f"  [{i}] 图4.5 → 图4.7 (个人中心界面)")
    elif t == '图4.6  视频编辑操作流程图':
        for run in p.runs:
            run.text = run.text.replace('图4.6', '图4.8')
        print(f"  [{i}] 图4.6 → 图4.8 (操作流程图)")

print("\n--- Fixing in-text references ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '图4.4展示了主页界面' in t:
        for run in p.runs:
            run.text = run.text.replace('图4.4', '图4.2')
        print(f"  [{i}] 图4.4→图4.2 (主页界面 ref fix)")
    elif '图4.5展示了项目列表界面' in t:
        for run in p.runs:
            run.text = run.text.replace('图4.5', '图4.3')
        print(f"  [{i}] 图4.5→图4.3 (项目列表 ref fix)")
    elif '图4.2展示了视频编辑器界面' in t:
        for run in p.runs:
            run.text = run.text.replace('图4.2展示了视频编辑器界面', '图4.6展示了视频编辑器界面')
        print(f"  [{i}] 图4.2→图4.6 (视频编辑器 ref renumber)")
    elif '图4.3展示了个人中心界面' in t:
        for run in p.runs:
            run.text = run.text.replace('图4.3展示了个人中心界面', '图4.7展示了个人中心界面')
        print(f"  [{i}] 图4.3→图4.7 (个人中心 ref renumber)")
    elif '图4.6展示了用户使用本系统' in t:
        for run in p.runs:
            run.text = run.text.replace('图4.6', '图4.8')
        print(f"  [{i}] 图4.6→图4.8 (操作流程 ref renumber)")

doc.save(out_path)
print(f"\nSaved to: {out_path}")

# Verification
print("\n--- Verification ---")
doc2 = docx.Document(out_path)
import re
print("Captions:")
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if re.match(r'^图4\.\d+', t):
        print(f"  [{i}] {t}")

print("In-text refs:")
for i, p in enumerate(doc2.paragraphs):
    t = p.text
    refs = re.findall(r'图4\.\d+', t)
    if refs and not re.match(r'^图4\.\d+', t.strip()):
        for ref in refs:
            idx = t.index(ref)
            print(f"  [{i}] {ref} | ...{t[max(0,idx-8):idx+25]}...")
