import docx
from docx.oxml.ns import qn

doc_path = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"
doc = docx.Document(doc_path)

print("=== 验证修改后的文档 ===\n")

# Check figures
print("--- 图片段落 ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    has_img = bool(para._element.findall('.//' + qn('wp:inline')))
    
    if has_img:
        next_text = doc.paragraphs[i+1].text.strip() if i < len(doc.paragraphs)-1 else ""
        print(f"P{i} [IMG] next: {next_text[:60]}")
    
    if text.startswith('图 3.') or text.startswith('图 4.'):
        print(f"P{i} CAPTION: {text}")
    
    if '（4-' in text or 'Clarity' in text or 'Score' in text or 'Motion' in text or 'Relevance' in text or 'Confidence' in text or 'Trend' in text:
        print(f"P{i} FORMULA: {text[:100].encode('gbk', errors='replace').decode('gbk')}")

print("\n--- 新增的后端流程图 ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '图 4.28' in text or '图 4.29' in text or '图 4.30' in text or '图 4.31' in text:
        prev_has_img = bool(doc.paragraphs[i-1]._element.findall('.//' + qn('wp:inline'))) if i > 0 else False
        print(f"P{i} {text} (prev_has_img={prev_has_img})")

print("\n--- 总段落数 ---")
print(f"原始文档段落数参考: ~401")
print(f"修改后段落数: {len(doc.paragraphs)}")
