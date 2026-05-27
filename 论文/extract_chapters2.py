import docx
from docx.oxml.ns import qn

doc_path = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20.docx"
doc = docx.Document(doc_path)

output_lines = []

# First pass: find all headings
output_lines.append("=== ALL HEADINGS ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style_name = para.style.name if para.style else "None"
    if 'Heading' in style_name or '标题' in style_name or style_name.startswith('heading'):
        output_lines.append(f"P{i} [{style_name}] {text}")

# Second pass: extract chapters 3 and 4 by paragraph index
output_lines.append("\n=== CHAPTER 3 & 4 FULL CONTENT ===\n")
in_target = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style_name = para.style.name if para.style else "None"
    
    # Start at 3.1 area (after 图 2.1)
    if i >= 185 and not in_target:
        in_target = True
    
    # Stop after chapter 5 starts
    if i > 345:
        in_target = False
    
    if in_target:
        has_drawing = bool(para._element.findall('.//' + qn('wp:inline')) or para._element.findall('.//' + qn('wp:anchor')))
        marker = "[IMG]" if has_drawing else ""
        if text or has_drawing:
            output_lines.append(f"P{i} [{style_name}] {marker} {text}")

result = "\n".join(output_lines)
with open(r"F:\26毕设2\移动端短视频智能剪辑app\论文\chapters_34.txt", "w", encoding="utf-8") as f:
    f.write(result)

print(f"Done. {len(output_lines)} lines.")
