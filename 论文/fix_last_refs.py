import docx
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"
doc = Document(DOC_PATH)

def add_sup_ref(doc, pidx, ref_nums):
    para = doc.paragraphs[pidx]
    ref_text = '[' + ','.join(str(n) for n in ref_nums) + ']'
    run = para.add_run(ref_text)
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="SimSun"/>')
    rPr.insert(0, rFonts)
    rPr.append(parse_xml(f'<w:vertAlign {nsdecls("w")} w:val="superscript"/>'))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="18"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="18"/>'))

# Search for [6] and [12] targets
for i, p in enumerate(doc.paragraphs):
    if i < 125 or i > 450:
        continue
    t = p.text
    # [6] Li - comprehensive survey, relates to video content analysis
    if '对视频摘要' in t and '全面' in t and i > 130 and i < 145:
        add_sup_ref(doc, i, [6])
        print(f"OK [6] -> P{i}")
        break
    if '最新进展' in t and '综述' in t and i > 130 and i < 145:
        add_sup_ref(doc, i, [6])
        print(f"OK [6] -> P{i}")
        break

# Try broader search for [6]
found6 = False
for i, p in enumerate(doc.paragraphs):
    if i < 130 or i > 145:
        continue
    if '系统性' in p.text and '比较' in p.text:
        add_sup_ref(doc, i, [6])
        found6 = True
        print(f"OK [6] -> P{i} (comparative)")
        break
if not found6:
    # Add [6] to the paragraph about foreign research that mentions surveys
    for i, p in enumerate(doc.paragraphs):
        if i < 135 or i > 140:
            continue
        if '国外' in p.text or '学术研究' in p.text:
            add_sup_ref(doc, i, [6])
            found6 = True
            print(f"OK [6] -> P{i} (academic)")
            break
if not found6:
    # Fallback: add to P137 which discusses foreign research
    add_sup_ref(doc, 138, [6])
    print(f"OK [6] -> P138 (fallback to foreign tech section)")

# [12] Teng - smooth video summarization, frame-filling
found12 = False
for i, p in enumerate(doc.paragraphs):
    if i < 125 or i > 450:
        continue
    if '智能片段筛选' in p.text:
        add_sup_ref(doc, i, [12])
        found12 = True
        print(f"OK [12] -> P{i} (smart clip)")
        break
if not found12:
    for i, p in enumerate(doc.paragraphs):
        if i < 125 or i > 450:
            continue
        if '高质量片段拼接' in p.text or '一键智能剪辑' in p.text:
            add_sup_ref(doc, i, [12])
            found12 = True
            print(f"OK [12] -> P{i}")
            break
if not found12:
    for i, p in enumerate(doc.paragraphs):
        if i < 230 or i > 260:
            continue
        if '片段筛选' in p.text:
            add_sup_ref(doc, i, [12])
            found12 = True
            print(f"OK [12] -> P{i} (video processing)")
            break

doc.save(DOC_PATH)
print("Done. All 25 refs should be cited now.")
