"""插入用例图到修改版文档，替换图2.1"""
from docx import Document
from docx.shared import Inches, Cm
import shutil
import os

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"
IMG_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\figures\fig21_usecase.png"

BAK_PATH = DOC_PATH + ".bak_before_usecase"
shutil.copy2(DOC_PATH, BAK_PATH)
print(f"备份完成: {BAK_PATH}")

doc = Document(DOC_PATH)

target_caption = "图 2.1"
target_para_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if target_caption in text and "用例图" in text:
        target_para_idx = i
        print(f"找到图2.1标题: P{i} -> '{text}'")
        break

if target_para_idx is None:
    print("ERROR: 未找到图2.1标题")
    exit(1)

img_para_idx = target_para_idx - 1
img_para = doc.paragraphs[img_para_idx]

has_drawing = False
for run in img_para.runs:
    if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or \
       run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        has_drawing = True
        break

if not has_drawing:
    from lxml import etree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}
    drawings = img_para._element.findall('.//wp:anchor', ns) + img_para._element.findall('.//wp:inline', ns)
    if drawings:
        has_drawing = True

print(f"图片段落 P{img_para_idx}: '{img_para.text[:50] if img_para.text else '(empty)'}', has_drawing={has_drawing}")

if has_drawing:
    for run in img_para.runs:
        run.clear()

    from lxml import etree
    for child in list(img_para._element):
        tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ''
        if tag in ('r', 'drawing', 'pict'):
            img_para._element.remove(child)

    run = img_para.add_run()
    run.add_picture(IMG_PATH, width=Cm(14))
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("已替换图2.1中的旧图片")
else:
    print(f"WARNING: P{img_para_idx} 没有图片，尝试搜索附近段落")
    for offset in [-2, -1, 0, 1, 2]:
        check_idx = target_para_idx + offset
        if 0 <= check_idx < len(doc.paragraphs):
            p = doc.paragraphs[check_idx]
            from lxml import etree
            ns = {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}
            drawings = p._element.findall('.//wp:anchor', ns) + p._element.findall('.//wp:inline', ns)
            print(f"  P{check_idx}: text='{p.text[:40] if p.text else '(empty)'}', drawings={len(drawings)}")

doc.save(DOC_PATH)
print(f"文档已保存: {DOC_PATH}")
