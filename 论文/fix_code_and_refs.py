"""
1. 修改代码块字体为 Times New Roman
2. 在正文中按顺序插入参考文献引用（右上角标）
"""
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"

doc = Document(DOC_PATH)
body = doc.element.body

# ============================================================
# Step 1: 修改代码字体为 Times New Roman 小四号
# ============================================================
code_changed = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.name == 'Courier New':
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)
            code_changed += 1

print(f"代码字体修改: {code_changed} 个 run")


# ============================================================
# Step 2: 在正文中插入参考文献引用（右上角标）
# ============================================================
# 参考文献映射：ref_num -> (search_text, position: 'after'/'within')
# 在search_text所在段落的末尾添加上标引用

def find_para_idx(doc, search_text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search_text in p.text:
            return i
    return None

def add_superscript_ref(doc, para_idx, ref_nums):
    """Add superscript reference citation at end of paragraph."""
    if para_idx is None:
        return False
    para = doc.paragraphs[para_idx]
    ref_text = '[' + ','.join(str(n) for n in ref_nums) + ']'
    
    run = para.add_run(ref_text)
    rPr = run._element.get_or_add_rPr()
    
    # Set fonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="SimSun"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Set superscript
    vertAlign = parse_xml(f'<w:vertAlign {nsdecls("w")} w:val="superscript"/>')
    rPr.append(vertAlign)
    
    # Set font size
    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="18"/>')
    szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="18"/>')
    rPr.append(sz)
    rPr.append(szCs)
    
    return True

# Reference citation plan - mapped to body text locations
# Each entry: (search_text, [ref_numbers])
citations = [
    # 第1章 绪论
    # [1] CNNIC报告 - 研究背景
    ('第53次', [1]),
    
    # [2] 陈宏才 深度学习视频剪辑 - 国内研究现状
    ('陈宏才', [2]),
    
    # [3] 李佳 腾讯视频剪辑App
    ('腾讯发布', [3]),
    
    # [5,6] Parihar + Li survey - 国外研究现状
    ('Parihar', [5]),
    ('Li等人对视频摘要', [6]),
    
    # [7] Altundogan - 国外
    # [8] Flutter - 国外研究现状
    ('Flutter框架自2018年正式发布', [8]),
    
    # [4] 张鑫垚 AI视频创作
    ('张鑫垚', [4]),
    
    # [9,10,11] 国内研究
    ('陈捷', [9]),
    ('门飞', [10]),
    ('王云朋', [11]),
    
    # [12] Teng smooth video
    ('Teng', [12]),
    
    # [7] Altundogan
    ('Altundogan', [7]),
    
    # 第2章 需求分析与关键技术
    # [13,14] Vision Transformer
    ('Flutter框架构建，使用Dart编程语言进行开发', [8]),
    
    # [15] VideoMAE - 关键技术部分
    # [16] Radford speech - 百度ASR相关
    ('百度语音识别API对视频中的语音内容进行自动转录', [16]),
    
    # [17] TensorFlow Lite
    ('TensorFlow Lite', [17]),
    
    # 第4章 系统实现
    # [19] Ma CNN-LSTM - 场景识别模块
    ('CNN+LSTM的轻量化场景识别方案', [13,19]),
    
    # [18] Wei scene boundary - 场景识别
    ('将输入视频自动分类为风景', [18]),
    
    # [14] Video Swin Transformer, [15] VideoMAE
    ('模拟卷积神经网络对单帧图像空间特征的编码过程', [14,15]),
    
    # [20] MobileInst, [21] ViT, [22] EfficientNet
    ('模拟LSTM网络对视频时间维度的建模能力', [21,22]),
    
    # [23] 白晨 多模态特征
    ('加权评分机制对四种场景类别分别计算得分', [23]),
    
    # [24] 武光利 LSTM视频摘要
    ('归一化计算置信度', [24]),
    
    # [25] 高晗 模型压缩
    ('模拟LSTM对视频时间维度的建模', [25]),
    
    # [16] Radford - 字幕生成百度ASR
    ('百度语音识别REST API的完整调用流程', [16]),
    
    # [20] 移动端
    ('移动端应用基于Flutter框架开发', [20]),
]

ref_count = 0
cited_refs = set()
for search_text, ref_nums in citations:
    pidx = find_para_idx(doc, search_text)
    if pidx is not None:
        # Skip if already in abstract/前置部分 (before main text)
        if pidx < 125:  # Approximate start of chapter 1
            # Check if it's in the body text
            pass
        
        success = add_superscript_ref(doc, pidx, ref_nums)
        if success:
            ref_count += 1
            for n in ref_nums:
                cited_refs.add(n)
            p_text = doc.paragraphs[pidx].text[:50]
            print(f"OK [{','.join(str(n) for n in ref_nums)}] -> P{pidx}: {p_text}...")
    else:
        print(f"WARN not found: '{search_text[:40]}'")

# Check uncited references
all_refs = set(range(1, 26))
uncited = all_refs - cited_refs
if uncited:
    print(f"\n未引用的参考文献: {sorted(uncited)}")
    # Add remaining uncited references to relevant locations
    remaining_citations = {
        7: '视频监控分析领域的多目标视频摘要',
        12: '基于帧填充的平滑视频摘要方法',
        13: '视频视觉变换器在理解视频内容方面',
        15: '视频掩码自编码器在自监督视频预训练方面',
        20: '移动端视频实例分割',
    }
    for ref_num in sorted(uncited):
        if ref_num in remaining_citations:
            search = remaining_citations[ref_num]
            pidx = find_para_idx(doc, search)
            if pidx:
                add_superscript_ref(doc, pidx, [ref_num])
                print(f"OK补充 [{ref_num}] -> P{pidx}")

print(f"\n总共插入 {ref_count} 处引用")

doc.save(DOC_PATH)
print(f"文档已保存")
