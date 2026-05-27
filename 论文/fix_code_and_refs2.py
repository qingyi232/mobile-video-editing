"""
1. 修改代码块字体为 Times New Roman
2. 在正文中按顺序插入参考文献引用（右上角标）
重新从原修改版文档开始操作。
"""
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import shutil

SRC = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20.docx"
DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"

# Reload from the insert_all output
doc = Document(DOC_PATH)

# ============================================================
# Step 1: 修改代码字体为 Times New Roman
# ============================================================
code_changed = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.name == 'Courier New':
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)
            code_changed += 1
print(f"Step1: code font changed: {code_changed} runs")

# ============================================================
# Step 2: 在正文中添加参考文献引用
# ============================================================
BODY_START = 125  # Chapter 1 starts approximately here
BODY_END = 420    # Before references section

def find_para_in_body(doc, search_text, start=BODY_START, end=BODY_END):
    for i, p in enumerate(doc.paragraphs):
        if i < start or i > end:
            continue
        if search_text in p.text:
            return i
    return None

def add_sup_ref(doc, pidx, ref_nums):
    if pidx is None:
        return False
    para = doc.paragraphs[pidx]
    ref_text = '[' + ','.join(str(n) for n in ref_nums) + ']'
    run = para.add_run(ref_text)
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="SimSun"/>')
    rPr.insert(0, rFonts)
    vertAlign = parse_xml(f'<w:vertAlign {nsdecls("w")} w:val="superscript"/>')
    rPr.append(vertAlign)
    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="18"/>')
    szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="18"/>')
    rPr.append(sz)
    rPr.append(szCs)
    return True

# Citation plan: (search_text_in_body, [ref_numbers])
# Must be in body text (Ch1-Ch5), not in abstract or reference list
citations = [
    # [1] CNNIC报告 - 1.1.1 研究背景
    ('中国互联网络信息中心', [1]),
    
    # [2] 陈宏才 - 1.2.2 国内研究
    ('基于深度学习的视频自动剪辑与拼接技术', [2]),
    
    # [3] 李佳 腾讯视频剪辑 - 1.2.2
    ('功能体系', [3]),
    
    # [4] 张鑫垚 AI视频创作 - 1.2.2
    ('情感感知与智能剪辑', [4]),
    
    # [5] Parihar survey - 1.2.1 国外
    ('视频摘要生成技术进行了系统性', [5]),
    
    # [6] Li comprehensive survey - 1.2.1
    ('视频摘要技术的挑战和最新进展进行了全面', [6]),
    
    # [7] Altundogan video surveillance - 1.2.1
    ('视频监控分析领域的多目标视频摘要', [7]),
    
    # [8] Flutter 2018 - 1.2.1 技术发展
    ('Flutter框架自2018年正式发布', [8]),
    
    # [9] 陈捷 - 1.2.2
    ('智能剪辑背景下AI技术在微视频', [9]),
    
    # [10] 门飞 - 1.2.2
    ('感知与智能剪辑技术在视频创作中的', [10]),
    
    # [11] 王云朋 - 1.2.2
    ('人工智能技术在广播影视剪辑中的', [11]),
    
    # [12] Teng frame-filling - 1.2.1
    ('基于帧填充的平滑视频摘要方法', [12]),
    
    # [13] Arnab ViViT - 第2章关键技术
    ('CNN+LSTM的轻量化场景识别方案', [13]),
    
    # [14] Liu Video Swin Transformer - 场景识别
    ('模拟卷积神经网络对单帧图像空间特征的编码过程', [14]),
    
    # [15] Tong VideoMAE
    ('将各帧的清晰度特征按时间顺序排列组成序列', [15]),
    
    # [16] Radford speech recognition - 字幕
    ('百度语音识别API对视频中的语音内容进行自动转录', [16]),
    
    # [17] TensorFlow Lite - 关键技术
    ('TensorFlow Lite技术的成熟', [17]),
    
    # [18] Wei scene boundary - 场景识别
    ('自动分类为风景', [18]),
    
    # [19] Ma CNN-LSTM - 场景识别
    ('CNN特征提取阶段', [19]),
    
    # [20] Shi MobileInst - 移动端
    ('移动端应用基于Flutter框架开发', [20]),
    
    # [21] Dosovitskiy ViT - 视觉变换器
    ('图像的边缘密度和纹理丰富度', [21]),
    
    # [22] Tan EfficientNetV2 - 高效网络
    ('模拟LSTM网络对视频时间维度的建模能力', [22]),
    
    # [23] 白晨 多模态特征 - 场景识别
    ('加权评分机制对四种场景类别分别计算得分', [23]),
    
    # [24] 武光利 双向LSTM - 场景识别
    ('归一化计算置信度', [24]),
    
    # [25] 高晗 深度学习压缩 - CNN+LSTM
    ('CNN+LSTM模型的视频场景自动识别', [25]),
]

ref_count = 0
cited_refs = set()

for search_text, ref_nums in citations:
    pidx = find_para_in_body(doc, search_text)
    if pidx is not None:
        success = add_sup_ref(doc, pidx, ref_nums)
        if success:
            ref_count += 1
            for n in ref_nums:
                cited_refs.add(n)
            safe_text = doc.paragraphs[pidx].text[:40].encode('gbk', errors='replace').decode('gbk')
            print(f"  OK [{','.join(str(n) for n in ref_nums)}] -> P{pidx}: {safe_text}...")
    else:
        print(f"  MISS: '{search_text[:40]}'")

all_refs = set(range(1, 26))
uncited = all_refs - cited_refs
print(f"\nStep2: {ref_count} citations added, refs cited: {sorted(cited_refs)}")
if uncited:
    print(f"  Uncited refs: {sorted(uncited)}")

doc.save(DOC_PATH)
print(f"\nSaved to: {DOC_PATH}")
