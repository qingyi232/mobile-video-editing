"""
将所有参考文献引用改为真正的Word交叉引用（Ctrl+点击可跳转）。
1. 先删除之前插入的纯文本上标引用
2. 给参考文献列表每条添加书签
3. 在正文中插入REF域代码交叉引用
"""
import docx
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
import re
import copy

DOC_PATH = r"F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx"

doc = Document(DOC_PATH)

# ============================================================
# Step 1: 删除之前插入的纯文本上标引用 [N]
# ============================================================
removed = 0
for para in doc.paragraphs:
    runs_to_remove = []
    for run in para.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            va = rPr.find(qn('w:vertAlign'))
            if va is not None and va.get(qn('w:val')) == 'superscript':
                if re.match(r'^\[\d+(,\d+)*\]$', run.text.strip()):
                    runs_to_remove.append(run)
    for run in runs_to_remove:
        run._element.getparent().remove(run._element)
        removed += 1

print(f"Step1: Removed {removed} old superscript citations")

# ============================================================
# Step 2: 找到参考文献列表，给每条添加书签
# ============================================================
BODY_START = 125
ref_section_start = None
ref_paras = {}  # {ref_num: para_idx}

for i, p in enumerate(doc.paragraphs):
    if i < BODY_START:
        continue
    text = p.text.strip()
    if '参考文献' in text and len(text) < 20:
        ref_section_start = i
        continue
    if ref_section_start and text.startswith('['):
        match = re.match(r'\[(\d+)\]', text)
        if match:
            num = int(match.group(1))
            if num not in ref_paras:
                ref_paras[num] = i

print(f"Step2: Found {len(ref_paras)} references: {sorted(ref_paras.keys())}")

# Add bookmarks to each reference paragraph
for ref_num, pidx in ref_paras.items():
    para = doc.paragraphs[pidx]
    bm_name = f"_Ref_{ref_num}"
    
    # Create bookmark start before first run
    bm_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{ref_num}" w:name="{bm_name}"/>'
    )
    bm_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{ref_num}"/>'
    )
    
    # Insert bookmark wrapping the paragraph content
    first_elem = para._element[0] if len(para._element) else None
    if first_elem is not None:
        para._element.insert(0, bm_start)
        para._element.append(bm_end)
    
print(f"Step2: Added {len(ref_paras)} bookmarks")

# ============================================================
# Step 3: 在正文中插入交叉引用
# ============================================================
def find_para_in_body(doc, search_text, start=BODY_START, end=None):
    if end is None:
        end = ref_section_start or 500
    for i, p in enumerate(doc.paragraphs):
        if i < start or i > end:
            continue
        if search_text in p.text:
            return i
    return None

def add_cross_ref(doc, pidx, ref_nums):
    """Add a Word cross-reference field to the end of a paragraph."""
    if pidx is None:
        return False
    
    para = doc.paragraphs[pidx]
    ref_text = '[' + ','.join(str(n) for n in ref_nums) + ']'
    
    # For each ref number, create a REF field with hyperlink
    for idx, ref_num in enumerate(ref_nums):
        bm_name = f"_Ref_{ref_num}"
        
        # Common rPr for superscript
        rpr_xml = (
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="SimSun"/>'
            f'<w:vertAlign w:val="superscript"/>'
            f'<w:sz w:val="18"/><w:szCs w:val="18"/>'
            f'</w:rPr>'
        )
        
        # Add "[" before first ref
        if idx == 0:
            r_bracket = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:t>[</w:t></w:r>')
            para._element.append(r_bracket)
        
        # Field begin
        r_begin = parse_xml(
            f'<w:r {nsdecls("w")}>{rpr_xml}'
            f'<w:fldChar w:fldCharType="begin"/></w:r>'
        )
        para._element.append(r_begin)
        
        # Field instruction
        r_instr = parse_xml(
            f'<w:r {nsdecls("w")}>{rpr_xml}'
            f'<w:instrText xml:space="preserve"> REF {bm_name} \\h </w:instrText></w:r>'
        )
        para._element.append(r_instr)
        
        # Field separate
        r_sep = parse_xml(
            f'<w:r {nsdecls("w")}>{rpr_xml}'
            f'<w:fldChar w:fldCharType="separate"/></w:r>'
        )
        para._element.append(r_sep)
        
        # Display text
        display = str(ref_num)
        r_text = parse_xml(
            f'<w:r {nsdecls("w")}>{rpr_xml}'
            f'<w:t>{display}</w:t></w:r>'
        )
        para._element.append(r_text)
        
        # Field end
        r_end = parse_xml(
            f'<w:r {nsdecls("w")}>{rpr_xml}'
            f'<w:fldChar w:fldCharType="end"/></w:r>'
        )
        para._element.append(r_end)
        
        # Add comma between multiple refs
        if idx < len(ref_nums) - 1:
            r_comma = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:t>,</w:t></w:r>')
            para._element.append(r_comma)
    
    # Add "]" after last ref
    r_bracket_end = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:t>]</w:t></w:r>')
    para._element.append(r_bracket_end)
    
    return True

# Same citation plan as before
citations = [
    ('中国互联网络信息中心', [1]),
    ('基于深度学习的视频自动剪辑与拼接技术', [2]),
    ('功能体系', [3]),
    ('情感感知与智能剪辑', [4]),
    ('视频摘要生成技术进行了系统性', [5]),
    ('学术研究方面', [6]),
    ('Adobe Premiere Pro', [7]),
    ('Flutter框架自2018年正式发布', [8]),
    ('智能剪辑背景下AI技术在微视频', [9]),
    ('感知与智能剪辑技术在视频创作中的', [10]),
    ('人工智能技术在广播影视剪辑中的', [11]),
    ('智能片段筛选', [12]),
    ('CNN+LSTM的轻量化场景识别方案', [13]),
    ('模拟卷积神经网络对单帧图像空间特征的编码过程', [14]),
    ('将各帧的清晰度特征按时间顺序排列组成序列', [15]),
    ('百度语音识别API对视频中的语音内容进行自动转录', [16]),
    ('CNN+LSTM模型', [17]),
    ('自动分类为风景', [18]),
    ('CNN特征提取阶段', [19]),
    ('移动端应用基于Flutter框架开发', [20]),
    ('图像的边缘密度和纹理丰富度', [21]),
    ('模拟LSTM网络对视频时间维度的建模能力', [22]),
    ('加权评分机制对四种场景类别分别计算得分', [23]),
    ('归一化计算置信度', [24]),
    ('CNN+LSTM模型的视频场景自动识别', [25]),
]

ref_count = 0
cited = set()

for search_text, ref_nums in citations:
    pidx = find_para_in_body(doc, search_text)
    if pidx is not None:
        # Verify all ref_nums have bookmarks
        valid_refs = [n for n in ref_nums if n in ref_paras]
        if valid_refs:
            success = add_cross_ref(doc, pidx, valid_refs)
            if success:
                ref_count += 1
                for n in valid_refs:
                    cited.add(n)
                safe = doc.paragraphs[pidx].text[:40].encode('gbk', errors='replace').decode('gbk')
                print(f"  OK [{','.join(str(n) for n in valid_refs)}] P{pidx}: {safe}...")
    else:
        print(f"  MISS: '{search_text[:30]}'")

uncited = set(range(1, 26)) - cited
print(f"\nStep3: {ref_count} cross-refs added, cited: {sorted(cited)}")
if uncited:
    print(f"  Uncited: {sorted(uncited)}")

doc.save(DOC_PATH)
print(f"\nSaved with cross-references.")
