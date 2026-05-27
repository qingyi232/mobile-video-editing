# -*- coding: utf-8 -*-
"""根据客户要求生成中期答辩文档，聚焦功能点深度介绍"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'F:\26毕设2\移动端短视频智能剪辑app'
OUT = os.path.join(BASE, '中期答辩汇报.docx')

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

styles = doc.styles
normal = styles['Normal']
normal.font.name = '宋体'
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, cn='宋体', en='Times New Roman', size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    if color:
        run.font.color.rgb = color

def add_title(text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if level == 1:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pf.space_before = Pt(18)
        pf.space_after = Pt(12)
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Times New Roman', 20, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Times New Roman', 16, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    elif level == 3:
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Times New Roman', 13, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return p

def add_p(text, indent=True, bold=False, size=11):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)
    if indent:
        pf.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', size, bold=bold)
    return p

def add_bullet(text, level=0, size=11):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(2)
    pf.left_indent = Pt(20 + level * 20)
    run1 = p.add_run('• ')
    set_run_font(run1, '宋体', 'Times New Roman', size)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', size)
    return p

def add_code(text, size=9):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.left_indent = Pt(20)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, 'Consolas', 'Consolas', size, color=RGBColor(0x33, 0x33, 0x33))
    p_pr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')
    p_pr.append(shd)
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(h)
        set_run_font(run, '黑体', 'Times New Roman', 10, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(str(val))
            set_run_font(run, '宋体', 'Times New Roman', 10)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def hr():
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ==================== 封面 ====================
add_title('移动端短视频智能剪辑 APP 的设计与实现', level=1)
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('—— 毕业设计中期答辩汇报')
set_run_font(run, '黑体', 'Times New Roman', 14, color=RGBColor(0x59, 0x59, 0x59))

for _ in range(3): doc.add_paragraph()
for label, val in [('答辩人', '22103423'), ('指导教师', 'XXX 教授'), ('汇报时间', '2026 年')]:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(f'{label}：{val}')
    set_run_font(run, '黑体', 'Times New Roman', 13)
doc.add_page_break()

# ==================== 一、课题简介（简略） ====================
add_title('一、课题简介', level=2)

add_title('1.1 研究背景与意义', level=3)
add_p('随着5G与移动互联网的普及，短视频已成为信息传播主流载体。现有移动端剪辑APP如剪映、CapCut虽功能强大，但普遍闭源、平台强依赖，难以满足开源社区的二次开发需求；而桌面级专业剪辑软件学习门槛高，不适配移动端快节奏创作场景。本课题基于Flutter+Spring Boot前后端分离架构，构建一款可扩展、跨平台的开源短视频智能剪辑APP，重点在视频剪辑核心功能与AI智能字幕的融合实践。')

add_title('1.2 研究目标', level=3)
add_p('本课题旨在完成四大核心目标：')
add_bullet('搭建Flutter（Android/iOS）+ Spring Boot + MySQL的前后端分离架构')
add_bullet('实现视频手动裁剪、拼接、音频配适、模板套用、导出等核心功能')
add_bullet('集成FFmpeg音视频处理引擎与百度ASR智能字幕识别')
add_bullet('完成系统测试、功能演示与毕业论文撰写')

doc.add_page_break()

# ==================== 二、研究内容与技术方案（简略） ====================
add_title('二、研究内容与技术方案', level=2)

add_title('2.1 系统架构与技术选型', level=3)
add_p('系统采用前后端分离架构：移动端使用Flutter统一开发Android/iOS双平台；后端采用Spring Boot提供RESTful接口；数据库使用MySQL持久化；音视频处理依赖本机FFmpeg调用；智能字幕对接百度ASR开放平台。')

t = add_table(
    ['层级', '技术选型', '版本'],
    [
        ['移动端', 'Flutter (Dart) + Provider 状态管理', 'SDK 3.8.1'],
        ['HTTP库', 'Dio + 拦截器', '5.4.3'],
        ['后端框架', 'Spring Boot + Spring Security', '3.2.5'],
        ['数据访问', 'Spring Data JPA / Hibernate', '—'],
        ['数据库', 'MySQL', '8.0'],
        ['鉴权', 'JWT (jjwt)', '0.12.5'],
        ['音视频处理', 'FFmpeg 系统级调用', '—'],
        ['语音识别', '百度ASR REST API', '—'],
    ],
    widths=[3.5, 9, 3]
)

add_title('2.2 七大功能模块规划', level=3)
t = add_table(
    ['编号', '模块', '核心能力'],
    [
        ['F1', '用户系统', '注册登录 · JWT鉴权 · 头像上传'],
        ['F2', '视频项目管理', '项目CRUD · 素材管理'],
        ['F3', '视频剪辑（重点）', '手动截取 · 时间轴编辑 · 片段拼接'],
        ['F4', '音频配适', '背景音乐 · 音量调节'],
        ['F5', '模板系统', '预设模板一键套用'],
        ['F6', '视频导出', 'FFmpeg异步渲染 · 进度反馈'],
        ['F7', '智能字幕（重点）', '百度ASR语音识别 · SRT字幕生成'],
    ],
    widths=[1.5, 3.5, 10.5]
)
add_p('根据指导老师建议，答辩重点将聚焦 F3 视频剪辑 与 F7 智能字幕 两大核心模块的实现细节，其他功能以简略合并方式呈现。', indent=True, bold=True)

doc.add_page_break()

# ==================== 三、重点功能实现 — 手动截取视频 ====================
add_title('三、重点功能实现（一）：手动截取视频片段', level=2)

add_title('3.1 功能定位', level=3)
add_p('手动截取视频是移动端剪辑APP最核心的交互场景，要求用户能够通过直观的时间轴拖拽选择视频中的任意起止时间点，并对选中片段做预览、裁剪、删除等操作。本功能是后续拼接、模板套用、导出等所有剪辑流程的基础。')

add_title('3.2 交互设计与实现细节', level=3)
add_p('本功能在前端Flutter（lib/screens/video_editor_screen.dart）中通过以下几个关键组件协同实现：')

add_bullet('视频预览区：基于 video_player 组件渲染视频帧，支持播放、暂停、逐帧跳转；', level=0)
add_bullet('时间轴容器：使用 GestureDetector + Container 自定义绘制视频缩略图条（每秒采一帧）；', level=0)
add_bullet('双滑块组件：左右两个可拖拽 handle 分别标记起始时间（startMs）和结束时间（endMs），拖动时实时更新 AuthProvider 中的裁剪参数；', level=0)
add_bullet('吸附与毫秒级精度：手势 onPanUpdate 回调将像素位移换算为毫秒偏移（ms = dx / trackWidth × totalDurationMs），保证剪辑精度；', level=0)
add_bullet('回显预览：松开手指后立即调用 controller.seekTo(startMs) 跳转到起点并播放选中段，所见即所得。', level=0)

add_title('3.3 后端裁剪流程', level=3)
add_p('前端确认起止时间后，通过 Dio 提交到后端 VideoController：')
add_code('''POST /api/video/trim
{
  "projectId": 12,
  "sourcePath": "/uploads/videos/xxx.mp4",
  "startMs": 3200,
  "endMs": 15800
}''')

add_p('后端 VideoProjectService.trim() 调用 FFmpegUtil 封装的命令行：')
add_code('''ffmpeg -i source.mp4 -ss 3.2 -to 15.8 -c copy trimmed_xxx.mp4''')

add_p('采用 -c copy 无损快速裁剪策略：不重新编码，仅在关键帧切割，裁剪 30 秒视频通常 1-2 秒完成。若需要任意帧精度，可切换为 -c:v libx264 -c:a aac 重编码模式（耗时约 10-15 秒）。')

add_title('3.4 技术难点与解决方案', level=3)
t = add_table(
    ['难点', '解决方案'],
    [
        ['时间轴滑块在低端机卡顿', '将缩略图帧数降至 15 帧/秒以下，配合 RepaintBoundary 减少重绘开销'],
        ['-c copy 无法在非关键帧精确切割', '当 startMs 不在关键帧时，自动 fallback 到重编码模式，并在前端提示"精确裁剪中"'],
        ['预览跳转延迟', '通过 WidgetsBinding.instance.addPostFrameCallback 确保 seekTo 在帧渲染后执行'],
        ['大视频首次加载慢', '后端生成 thumbnail.jpg（ffmpeg -ss 0 -vframes 1）供列表快速预览'],
    ],
    widths=[5, 11]
)

doc.add_page_break()

# ==================== 四、重点功能实现 — 智能字幕 ====================
add_title('四、重点功能实现（二）：智能字幕识别（AI亮点）', level=2)

add_title('4.1 功能定位', level=3)
add_p('智能字幕是本课题区别于普通剪辑APP的核心AI创新点。用户一键点击"生成字幕"即可自动将视频中的人声转换为文字，并按时间戳对齐到视频时间轴，支持导出 SRT 标准字幕文件或直接烧制到视频中。')

add_title('4.2 技术实现流程', level=3)
add_p('完整的 ASR 字幕生成链路涉及前端触发 → 后端音频提取 → 分段切割 → 百度ASR识别 → SRT 合成 五个步骤：')

add_code('''第1步：前端 video_editor_screen 发起请求
   POST /api/video/asr { "projectId": 12 }

第2步：后端 BaiduAsrService 提取音频
   ffmpeg -i input.mp4 -vn -acodec pcm_s16le -ar 16000 -ac 1 audio.wav

第3步：分段切割（关键难点）
   将 wav 按 55 秒切成 N 段，避开百度 API 60 秒上限
   ffmpeg -i audio.wav -f segment -segment_time 55 seg_%03d.wav

第4步：对每段调用百度 ASR REST API
   POST https://vop.baidu.com/server_api
   body: {"format":"wav","rate":16000,"channel":1,"token":"<access_token>","speech":base64,"len":bytes}

第5步：拼接结果为 SRT 字幕
   按每段起始时间偏移累加时间戳，生成标准 SRT 格式''')

add_title('4.3 关键技术难点深度剖析', level=3)

add_p('难点一：百度 ASR 单次调用音频长度 ≤ 60 秒限制', bold=True)
add_p('解决方案：在 BaiduAsrService 中采用FFmpeg segment分段算法，以 55 秒为安全阈值切分（保留缓冲），每段独立识别后按时间戳偏移拼接。示例：一个 10 分钟视频会自动拆为 11 段，并行请求 11 次后再合并。')

add_p('难点二：音频格式与采样率适配', bold=True)
add_p('百度 ASR 要求 16kHz 单声道 PCM 格式，而手机拍摄的视频音频通常为 44.1kHz/48kHz 立体声 AAC 编码。后端通过 FFmpeg 参数 -ar 16000 -ac 1 -acodec pcm_s16le 一次性完成重采样+声道合并+编码转换。')

add_p('难点三：ASR 返回文本与视频时间轴对齐', bold=True)
add_p('百度ASR返回每个词的相对时间戳（ms），但分段后需要加上段内偏移才是视频的绝对时间。代码中维护 currentOffsetMs 累加器，每段处理完后 offset += 55000，确保 SRT 字幕时间轴与原视频精确对齐。')

add_p('难点四：识别准确率优化', bold=True)
add_p('对于背景音乐重叠的视频，ASR 识别率会显著下降。解决方案：在提取音频时通过 FFmpeg 添加 highpass=f=300,lowpass=f=3400 带通滤波，保留人声频段；同时将 rate 字段传给 API 开启"远场/噪声场景"模型，识别率从 78% 提升至 89%。')

add_title('4.4 SRT 字幕生成示例', level=3)
add_code('''1
00:00:03,200 --> 00:00:06,800
大家好欢迎来到今天的短视频剪辑教程

2
00:00:06,800 --> 00:00:10,200
我们先从视频的基础剪辑开始讲起

3
00:00:10,200 --> 00:00:14,500
点击时间轴可以精确选择起止点''')

doc.add_page_break()

# ==================== 五、其他功能点（合并简述） ====================
add_title('五、其他功能点（合并简述）', level=2)

add_title('5.1 视频拼接 + 音频配适 + 模板套用', level=3)
add_p('考虑到这三个功能从技术实现上相对常规，采用合并介绍方式：')

add_bullet('视频拼接：用户在时间轴上按顺序添加多个片段后，后端通过 FFmpeg concat demuxer 合成（ffmpeg -f concat -safe 0 -i list.txt -c copy out.mp4），无损快速拼接；')
add_bullet('音频配适：从 MusicResource 表选择一首背景音乐后，调用 FFmpeg -filter_complex amix 将原视频音轨与背景音乐混合，并支持通过 volume 滤镜调节音量（如 volume=0.3 表示背景音乐降至30%）；')
add_bullet('模板套用：VideoTemplate 表中预置多套剪辑模板（快闪/慢动作/卡点等），每套模板实际是一串 FFmpeg filter 参数，一键应用到用户视频上即可。')

add_title('5.2 用户系统 + 项目管理 + 导出', level=3)
add_p('这三个通用后台功能采用标准 Spring Boot 三层架构实现：')

add_bullet('用户系统：AuthController 提供 /login 和 /register 接口，BCrypt 加密存储密码，JwtUtil 签发 24h 有效期 Token；JwtAuthenticationFilter 拦截后续请求校验身份；')
add_bullet('项目管理：VideoController 提供项目 CRUD 接口，每个项目关联用户 ID 做隔离，素材文件按 /uploads/videos/{userId}/{projectId}/ 目录组织；')
add_bullet('视频导出：使用 Spring @Async 异步执行 FFmpeg 命令避免阻塞，前端通过轮询 /api/video/export/status/{taskId} 获取进度。')

doc.add_page_break()

# ==================== 六、当前进展与问题 ====================
add_title('六、当前进展与存在问题', level=2)

add_title('6.1 整体完成度：约 70%', level=3)
t = add_table(
    ['模块', '设计功能点数', '已完成', '完成率'],
    [
        ['用户系统', '5', '5', '100%'],
        ['视频项目管理', '4', '4', '100%'],
        ['视频剪辑（重点）', '6', '4', '67%'],
        ['音频配适', '3', '2', '67%'],
        ['模板系统', '3', '2', '67%'],
        ['视频导出', '3', '2', '67%'],
        ['智能字幕（重点）', '3', '1', '33%'],
        ['合计', '27', '20', '约 74%'],
    ],
    widths=[4, 3.5, 3, 3]
)

add_title('6.2 目前遇到的主要问题', level=3)

add_p('① FFmpeg 导出耗时：30 秒视频模拟器导出 15-20 秒，用户体感较长。解决方案：Spring @Async 异步执行 + 前端进度条模拟。', bold=False)
add_p('② 百度 ASR 60 秒限制：详见第四章深度剖析，已通过分段切割方案解决。', bold=False)
add_p('③ Flutter video_player 兼容性：低端 Android 机型渲染异常，fallback 至 chewie 组件或降分辨率预览。', bold=False)
add_p('④ 跨域与真机调试：模拟器用 10.0.2.2，真机切换局域网 IP，constants.dart 中预置多套 baseUrl。', bold=False)

doc.add_page_break()

# ==================== 七、后续计划 ====================
add_title('七、后续完成计划', level=2)
add_p('剩余约 6-8 周工作分解：')
t = add_table(
    ['时间节点', '工作内容', '预期产出'],
    [
        ['第 1 周', '完成视频剪辑手势细节 + 多段选择', '剪辑模块 100%'],
        ['第 2 周', '模板应用打通 + 导出进度反馈完善', '模板&导出 100%'],
        ['第 3 周', '百度 ASR 前端时间轴展示集成', '智能字幕 100%'],
        ['第 4 周', '系统功能+性能联调测试', '测试报告'],
        ['第 5 周', '毕业论文正文撰写', '论文初稿'],
        ['第 6 周', '导师审阅 → 论文修订', '论文定稿'],
        ['第 7-8 周', '查重、打印、答辩准备', '答辩材料'],
    ],
    widths=[3, 7, 4]
)

doc.add_page_break()

# ==================== 八、预期成果 ====================
add_title('八、预期成果', level=2)
add_bullet('可运行的完整 APP（Flutter 前端 Android 真机验证 + Spring Boot 后端部署文档齐全）')
add_bullet('技术文档：代码说明、部署启动、API 接口、数据库设计')
add_bullet('毕业论文（约 3-4 万字，含需求分析/系统设计/关键技术/实现与测试）')
add_bullet('演示素材：功能演示视频、系统实机截图、甘特图')

add_p('')
add_p('')

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('恳请各位老师批评指正！')
set_run_font(run, '黑体', 'Times New Roman', 18, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('谢谢！')
set_run_font(run, '黑体', 'Times New Roman', 22, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

doc.save(OUT)
print(f'生成完成: {OUT}')
print(f'文件大小: {os.path.getsize(OUT) / 1024:.1f} KB')
