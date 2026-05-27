# -*- coding: utf-8 -*-
"""
向论文 4.2.6 / 4.2.7 / 4.2.3 章节插入核心代码。
按从后到前的顺序插入，避免段落索引偏移问题。
"""
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
import copy

DOC_PATH = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx'
OUT_PATH = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版_已补代码.docx'

doc = Document(DOC_PATH)

CODE_TEMPLATE_REF = 230

CODE_426_CONTROLLER = '''@RestController
@RequestMapping("/api/templates")
@RequiredArgsConstructor
public class TemplateController {
    private final TemplateService templateService;

    @GetMapping
    public ApiResponse<List<VideoTemplate>> getAllTemplates() {
        return ApiResponse.success(
            templateService.getAllTemplates());
    }

    @GetMapping("/category/{category}")
    public ApiResponse<List<VideoTemplate>>
            getTemplatesByCategory(
            @PathVariable String category) {
        return ApiResponse.success(
            templateService.getTemplatesByCategory(category));
    }

    @GetMapping("/{id}")
    public ApiResponse<VideoTemplate> getTemplateById(
            @PathVariable Long id) {
        return ApiResponse.success(
            templateService.getTemplateById(id));
    }

    @PostMapping("/{id}/use")
    public ApiResponse<VideoTemplate> useTemplate(
            @PathVariable Long id) {
        return ApiResponse.success(
            templateService.useTemplate(id));
    }
}'''

CODE_426_SERVICE = '''@Service
@RequiredArgsConstructor
public class TemplateService {
    private final VideoTemplateRepository templateRepository;

    public List<VideoTemplate> getAllTemplates() {
        return templateRepository
            .findAllByOrderByUsageCountDesc();
    }

    public List<VideoTemplate> getTemplatesByCategory(
            String category) {
        return templateRepository.findByCategory(category);
    }

    public VideoTemplate getTemplateById(Long id) {
        return templateRepository.findById(id)
            .orElseThrow(() ->
                new RuntimeException("模板不存在"));
    }

    public VideoTemplate useTemplate(Long id) {
        VideoTemplate template = getTemplateById(id);
        template.setUsageCount(
            template.getUsageCount() + 1);
        return templateRepository.save(template);
    }
}'''

CODE_427_FILTER = '''public String applyFilter(String videoPath,
        String filterType, Long userId) {
    String absVideo = toAbsolutePath(videoPath);
    String dir = exportDir + "/" + userId;
    new File(dir).mkdirs();
    String outputPath = dir + "/filter_"
        + UUID.randomUUID() + ".mp4";
    boolean ok = ffmpegUtil.applyFilter(
        absVideo, outputPath, filterType);
    if (!ok)
        throw new RuntimeException("滤镜应用失败");
    validateOutput(outputPath, "滤镜应用");
    return toRelativePath(outputPath);
}'''

CODE_427_RATIO = '''public String changeAspectRatio(String videoPath,
        String ratio, Long userId) {
    String absVideo = toAbsolutePath(videoPath);
    String dir = exportDir + "/" + userId;
    new File(dir).mkdirs();
    String outputPath = dir + "/ratio_"
        + UUID.randomUUID() + ".mp4";
    boolean success = ffmpegUtil.changeAspectRatio(
        absVideo, outputPath, ratio);
    if (!success)
        throw new RuntimeException("比例转换失败");
    validateOutput(outputPath, "比例转换");
    return toRelativePath(outputPath);
}'''

CODE_423_CONTROLLER = '''@RestController
@RequestMapping("/api/music")
@RequiredArgsConstructor
public class MusicController {
    private final MusicService musicService;

    @GetMapping
    public ApiResponse<List<MusicResource>> getAllMusic() {
        return ApiResponse.success(
            musicService.getAllMusic());
    }

    @GetMapping("/category/{category}")
    public ApiResponse<List<MusicResource>>
            getMusicByCategory(
            @PathVariable String category) {
        return ApiResponse.success(
            musicService.getMusicByCategory(category));
    }

    @PostMapping("/recommend")
    public ApiResponse<List<MusicResource>>
            recommendMusic(
            @RequestBody MusicRecommendRequest request) {
        return ApiResponse.success(
            musicService.recommendMusic(request));
    }
}'''

CODE_423_SERVICE = '''public List<MusicResource> recommendMusic(
        MusicRecommendRequest request) {
    List<MusicResource> candidates = new ArrayList<>();
    String sceneType = request.getSceneType();
    if (sceneType != null) {
        sceneType = switch (sceneType) {
            case "calm" -> "nature";
            case "festive" -> "festive";
            default -> sceneType;
        };
    }
    if (sceneType != null
            && request.getMood() != null) {
        candidates = musicRepository
            .findByCategoryAndMood(
                sceneType, request.getMood());
    } else if (sceneType != null) {
        candidates = musicRepository
            .findByCategory(sceneType);
    } else if (request.getMood() != null) {
        candidates = musicRepository
            .findByMood(request.getMood());
    } else {
        candidates = musicRepository.findAll();
    }
    if (request.getVideoDuration() != null
            && request.getVideoDuration() > 0) {
        double targetDuration =
            request.getVideoDuration();
        candidates = candidates.stream()
            .sorted(Comparator.comparingDouble(m ->
                Math.abs(m.getDuration()
                    - targetDuration)))
            .collect(Collectors.toList());
    }
    return candidates.stream().limit(10)
        .collect(Collectors.toList());
}'''


def make_code_paragraph(template_p, line_text):
    new_p = copy.deepcopy(template_p._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    run_elem = template_p.runs[0]._element
    new_run = copy.deepcopy(run_elem)
    for t in new_run.findall(qn('w:t')):
        t.text = line_text if line_text.strip() else ' '
        t.set(qn('xml:space'), 'preserve')
    new_p.append(new_run)
    return new_p


def make_desc_paragraph(template_p, desc_text):
    new_p = copy.deepcopy(template_p._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    run_elem = template_p.runs[0]._element
    new_run = copy.deepcopy(run_elem)
    for t in new_run.findall(qn('w:t')):
        t.text = desc_text
        t.set(qn('xml:space'), 'preserve')
    new_p.append(new_run)
    return new_p


def insert_code_block(doc, after_idx, desc_text, code_text, code_template_idx, desc_template_idx):
    template_code_p = doc.paragraphs[code_template_idx]
    template_desc_p = doc.paragraphs[desc_template_idx]

    ref_elem = doc.paragraphs[after_idx]._element

    desc_p = make_desc_paragraph(template_desc_p, desc_text)
    ref_elem.addnext(desc_p)

    lines = code_text.strip().split('\n')
    prev = desc_p
    for line in lines:
        code_p = make_code_paragraph(template_code_p, line)
        prev.addnext(code_p)
        prev = code_p

    return 1 + len(lines)


total_offset = 0

print("=== Step 1: 4.2.7 滤镜处理与画面比例调整 (after paragraph 350) ===")
n = insert_code_block(doc, 350 + total_offset, "滤镜处理核心代码如下：", CODE_427_FILTER, CODE_TEMPLATE_REF, 349)
total_offset += n
print(f"  Inserted {n} paragraphs (filter)")
n = insert_code_block(doc, 350 + total_offset, "画面比例调整核心代码如下：", CODE_427_RATIO, CODE_TEMPLATE_REF, 349)
total_offset += n
print(f"  Inserted {n} paragraphs (ratio)")

print(f"\n=== Step 2: 4.2.6 模板管理 (after paragraph 347) ===")
n = insert_code_block(doc, 347, "模板管理控制器核心代码如下：", CODE_426_CONTROLLER, CODE_TEMPLATE_REF, 346)
total_offset += n
print(f"  Inserted {n} paragraphs (controller)")
n = insert_code_block(doc, 347 + n, "模板服务层核心代码如下：", CODE_426_SERVICE, CODE_TEMPLATE_REF, 346)
total_offset += n
print(f"  Inserted {n} paragraphs (service)")

print(f"\n=== Step 3: 4.2.3 音乐推荐 (after paragraph 269) ===")
n = insert_code_block(doc, 269, "音乐推荐控制器核心代码如下：", CODE_423_CONTROLLER, CODE_TEMPLATE_REF, 267)
total_offset += n
print(f"  Inserted {n} paragraphs (controller)")
n = insert_code_block(doc, 269 + n, "音乐推荐服务层核心代码如下：", CODE_423_SERVICE, CODE_TEMPLATE_REF, 267)
total_offset += n
print(f"  Inserted {n} paragraphs (service)")

print(f"\nTotal paragraphs added: {total_offset}")
doc.save(OUT_PATH)
print(f"Saved to: {OUT_PATH}")
