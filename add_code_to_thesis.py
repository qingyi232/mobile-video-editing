# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
import copy

doc = Document(r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版.docx')

CODE_426 = '''@RestController
@RequestMapping("/api/templates")
@RequiredArgsConstructor
public class TemplateController {
    private final TemplateService templateService;

    @GetMapping
    public ApiResponse<List<VideoTemplate>> getAllTemplates() {
        return ApiResponse.success(templateService.getAllTemplates());
    }

    @GetMapping("/category/{category}")
    public ApiResponse<List<VideoTemplate>> getTemplatesByCategory(
            @PathVariable String category) {
        return ApiResponse.success(
            templateService.getTemplatesByCategory(category));
    }

    @GetMapping("/{id}")
    public ApiResponse<VideoTemplate> getTemplateById(@PathVariable Long id) {
        return ApiResponse.success(templateService.getTemplateById(id));
    }

    @PostMapping("/{id}/use")
    public ApiResponse<VideoTemplate> useTemplate(@PathVariable Long id) {
        return ApiResponse.success(templateService.useTemplate(id));
    }
}'''

CODE_426_SERVICE = '''@Service
@RequiredArgsConstructor
public class TemplateService {
    private final VideoTemplateRepository templateRepository;

    public List<VideoTemplate> getAllTemplates() {
        return templateRepository.findAllByOrderByUsageCountDesc();
    }

    public List<VideoTemplate> getTemplatesByCategory(String category) {
        return templateRepository.findByCategory(category);
    }

    public VideoTemplate getTemplateById(Long id) {
        return templateRepository.findById(id)
                .orElseThrow(() -> new RuntimeException("模板不存在"));
    }

    public VideoTemplate useTemplate(Long id) {
        VideoTemplate template = getTemplateById(id);
        template.setUsageCount(template.getUsageCount() + 1);
        return templateRepository.save(template);
    }
}'''

CODE_427_FILTER = '''public String applyFilter(String videoPath, String filterType,
                          Long userId) {
    String absVideo = toAbsolutePath(videoPath);
    String dir = exportDir + "/" + userId;
    new File(dir).mkdirs();
    String outputPath = dir + "/filter_"
        + UUID.randomUUID() + ".mp4";
    boolean ok = ffmpegUtil.applyFilter(
        absVideo, outputPath, filterType);
    if (!ok) throw new RuntimeException("滤镜应用失败");
    validateOutput(outputPath, "滤镜应用");
    return toRelativePath(outputPath);
}'''

CODE_427_RATIO = '''public String changeAspectRatio(String videoPath,
                               String ratio, Long userId) {
    String absVideo = toAbsolutePath(videoPath);
    String dir = exportDir + "/" + userId;
    new File(dir).mkdirs();
    String outputPath = dir + "/ratio_"
        + UUID.randomUUID() + ".mp4";
    boolean success = ffmpegUtil.changeAspectRatio(
        absVideo, outputPath, ratio);
    if (!success)
        throw new RuntimeException("比例转换失败");
    validateOutput(outputPath, "比例转换");
    return toRelativePath(outputPath);
}'''

CODE_423_CONTROLLER = '''@RestController
@RequestMapping("/api/music")
@RequiredArgsConstructor
public class MusicController {
    private final MusicService musicService;

    @GetMapping
    public ApiResponse<List<MusicResource>> getAllMusic() {
        return ApiResponse.success(musicService.getAllMusic());
    }

    @GetMapping("/category/{category}")
    public ApiResponse<List<MusicResource>> getMusicByCategory(
            @PathVariable String category) {
        return ApiResponse.success(
            musicService.getMusicByCategory(category));
    }

    @PostMapping("/recommend")
    public ApiResponse<List<MusicResource>> recommendMusic(
            @RequestBody MusicRecommendRequest request) {
        return ApiResponse.success(
            musicService.recommendMusic(request));
    }
}'''

CODE_423_SERVICE = '''public List<MusicResource> recommendMusic(
        MusicRecommendRequest request) {
    List<MusicResource> candidates = new ArrayList<>();
    String sceneType = request.getSceneType();
    if (sceneType != null) {
        sceneType = switch (sceneType) {
            case "calm" -> "nature";
            case "festive" -> "festive";
            default -> sceneType;
        };
    }
    if (sceneType != null && request.getMood() != null) {
        candidates = musicRepository
            .findByCategoryAndMood(sceneType, request.getMood());
    } else if (sceneType != null) {
        candidates = musicRepository.findByCategory(sceneType);
    } else if (request.getMood() != null) {
        candidates = musicRepository.findByMood(request.getMood());
    } else {
        candidates = musicRepository.findAll();
    }
    if (request.getVideoDuration() != null
            && request.getVideoDuration() > 0) {
        double targetDuration = request.getVideoDuration();
        candidates = candidates.stream()
            .sorted(Comparator.comparingDouble(m ->
                Math.abs(m.getDuration() - targetDuration)))
            .collect(Collectors.toList());
    }
    return candidates.stream().limit(10)
        .collect(Collectors.toList());
}'''

def add_code_paragraph(doc, index, code_text):
    template_p = doc.paragraphs[230]
    lines = code_text.strip().split('\n')
    inserted = 0
    for line in lines:
        new_p = copy.deepcopy(template_p._element)
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        run_elem = template_p.runs[0]._element
        new_run = copy.deepcopy(run_elem)
        for t in new_run.findall(qn('w:t')):
            t.text = line if line else ' '
            t.set(qn('xml:space'), 'preserve')
        new_p.append(new_run)
        ref_elem = doc.paragraphs[index + inserted]._element
        ref_elem.addnext(new_p)
        inserted += 1
    return inserted

insert_points = [
    (347, "\n模板管理控制器核心代码如下：", CODE_426),
    (None, "\n模板服务层核心代码如下：", CODE_426_SERVICE),
    (350 + 100, "\n滤镜处理核心代码如下：", CODE_427_FILTER),
    (None, "\n画面比例调整核心代码如下：", CODE_427_RATIO),
    (269 + 100, "\n音乐推荐控制器核心代码如下：", CODE_423_CONTROLLER),
    (None, "\n音乐推荐服务层核心代码如下：", CODE_423_SERVICE),
]

offset = 0

for idx_target, desc_text, code_str in [
    (347, "模板管理控制器核心代码如下：", CODE_426),
    (None, "模板服务层核心代码如下：", CODE_426_SERVICE),
]:
    if idx_target is not None:
        real_idx = idx_target + offset
    else:
        real_idx = last_idx

    desc_p = copy.deepcopy(doc.paragraphs[346]._element)
    for r in desc_p.findall(qn('w:r')):
        desc_p.remove(r)
    run_elem = doc.paragraphs[346].runs[0]._element
    new_run = copy.deepcopy(run_elem)
    for t in new_run.findall(qn('w:t')):
        t.text = desc_text
    desc_p.append(new_run)
    doc.paragraphs[real_idx]._element.addnext(desc_p)
    offset += 1
    real_idx += 1

    cnt = add_code_paragraph(doc, real_idx, code_str)
    offset += cnt
    last_idx = real_idx + cnt

real_350 = 350 + offset
for desc_text, code_str in [
    ("滤镜处理核心代码如下：", CODE_427_FILTER),
    ("画面比例调整核心代码如下：", CODE_427_RATIO),
]:
    desc_p = copy.deepcopy(doc.paragraphs[346]._element)
    for r in desc_p.findall(qn('w:r')):
        desc_p.remove(r)
    run_elem = doc.paragraphs[346].runs[0]._element
    new_run = copy.deepcopy(run_elem)
    for t in new_run.findall(qn('w:t')):
        t.text = desc_text
    desc_p.append(new_run)
    doc.paragraphs[real_350]._element.addnext(desc_p)
    offset += 1
    real_350 += 1
    cnt = add_code_paragraph(doc, real_350, code_str)
    offset += cnt
    real_350 += cnt

real_269 = 269 + offset
for desc_text, code_str in [
    ("音乐推荐控制器核心代码如下：", CODE_423_CONTROLLER),
    ("音乐推荐服务层核心代码如下：", CODE_423_SERVICE),
]:
    desc_p = copy.deepcopy(doc.paragraphs[267]._element)
    for r in desc_p.findall(qn('w:r')):
        desc_p.remove(r)
    run_elem = doc.paragraphs[267].runs[0]._element
    new_run = copy.deepcopy(run_elem)
    for t in new_run.findall(qn('w:t')):
        t.text = desc_text
    desc_p.append(new_run)
    doc.paragraphs[real_269]._element.addnext(desc_p)
    offset += 1
    real_269 += 1
    cnt = add_code_paragraph(doc, real_269, code_str)
    offset += cnt
    real_269 += cnt

output_path = r'F:\26毕设2\移动端短视频智能剪辑app\论文\22103423郭湘_移动端短视频智能剪辑APP的设计与实现_毕业设计说明书5.20_修改版_已补代码.docx'
doc.save(output_path)
print(f"Done! Saved to: {output_path}")
